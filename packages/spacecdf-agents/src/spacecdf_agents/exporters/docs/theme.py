"""SpaceCDF — Course-style DOCX theme.

Applies the uOttawa SEDTI / SpaceCDF Facilitator's Book visual identity to
python-docx documents.  Importable from both the agents-package design-review
generator and the server-side DID renderer.

Brand source: ``docs/assets/figures/uottawa_brand.py``.

Key entry points:
    - :func:`apply_styles` — register Heading/Body/Caption styles on a fresh Document.
    - :func:`add_cover_page` — slab-bold title block with crimson banner.
    - :func:`add_doc_info_table` — change record / document info matrix.
    - :func:`add_page_furniture` — bilingual EN/FR footer ("Page X of / de Y"),
      running header with crimson rule.
    - :func:`add_toc_field` — Word-native ToC field (updates on open).
    - :func:`add_section_break` — page break with continuation of furniture.
    - :func:`bookmarked_heading` — Heading + bookmark for cross-references.
    - :func:`styled_table` — Light-grid table with crimson header shading.
    - :func:`add_figure` — Embed a PNG with crimson "Figure N — caption" line.

All helpers are idempotent and safe to call multiple times on the same doc.
"""
from __future__ import annotations

from typing import Any

from docx import Document as _NewDocument
from docx.document import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import Table, _Cell

# ---------------------------------------------------------------------------
# Palette — uOttawa Horizon
# ---------------------------------------------------------------------------

GARNET = "8f001a"          # Primary garnet (uOttawa)
GARNET_2 = "9c1c30"        # Secondary garnet
GARNET_DARK = "5a0010"     # Deep garnet for ruled lines
CHARCOAL = "2d2d2c"        # Body text
CHARCOAL_2 = "3a3a37"
WARM_GREY = "80746c"       # Quiet metadata
WARM_GREY_2 = "908681"
BLUE = "636d77"            # Series accent 1
GREEN = "67796c"           # Series accent 2
POLAR = "f2f2f2"           # Soft fill
WHITE = "ffffff"

# Status colours kept compatible with legacy STATUS_COLOURS for callers.
STATUS_COLOURS = {
    "compliant": "C6EFCE", "green": "C6EFCE", "ok": "C6EFCE",
    "marginal": "FFEB9C", "amber": "FFEB9C", "warning": "FFEB9C",
    "non_compliant": "FFC7CE", "red": "FFC7CE", "exceeded": "FFC7CE",
    "violated": "FFC7CE", "fail": "FFC7CE",
}


def _rgb(hex6: str) -> RGBColor:
    return RGBColor.from_string(hex6.lstrip("#").upper())


# ---------------------------------------------------------------------------
# Document creation & global styles
# ---------------------------------------------------------------------------

def new_document() -> Document:
    """Create a fresh Document with the SpaceCDF page geometry + styles."""
    doc = _NewDocument()
    _set_page_geometry(doc)
    apply_styles(doc)
    return doc


def _set_page_geometry(doc: Document) -> None:
    # A4 = 21.0 cm.  Left + right margins of 1.8 cm give 17.4 cm usable text
    # width — comfortably above the widest table the renderer emits and
    # leaving a safety margin for cell padding.
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)


# Maximum table width that comfortably fits within the page margins
# (matches _set_page_geometry above minus a small safety margin).
USABLE_WIDTH_CM = 17.0


def apply_styles(doc: Document) -> None:
    """Register/override all SpaceCDF styles on the document."""
    styles = doc.styles

    # ---- Body text -------------------------------------------------------
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(CHARCOAL)
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15

    # ---- Headings --------------------------------------------------------
    _config_heading(styles, "Heading 1", size=20, bold=True, colour=GARNET,
                    space_before=Pt(20), space_after=Pt(8))
    _config_heading(styles, "Heading 2", size=15, bold=True, colour=GARNET,
                    space_before=Pt(14), space_after=Pt(4))
    _config_heading(styles, "Heading 3", size=12.5, bold=True, colour=CHARCOAL,
                    space_before=Pt(10), space_after=Pt(2))
    _config_heading(styles, "Heading 4", size=11, bold=True, colour=CHARCOAL_2,
                    space_before=Pt(8), space_after=Pt(2))

    # ---- Title (cover) ---------------------------------------------------
    try:
        title = styles["Title"]
    except KeyError:
        title = styles.add_style("Title", 1)  # 1 = paragraph
    title.font.name = "Calibri"
    title.font.size = Pt(36)
    title.font.bold = True
    title.font.color.rgb = _rgb(GARNET)
    title.paragraph_format.space_after = Pt(0)
    title.paragraph_format.space_before = Pt(0)

    # ---- Subtitle --------------------------------------------------------
    try:
        sub = styles["Subtitle"]
    except KeyError:
        sub = styles.add_style("Subtitle", 1)
    sub.font.name = "Cambria"
    sub.font.size = Pt(16)
    sub.font.italic = True
    sub.font.color.rgb = _rgb(CHARCOAL)

    # ---- Caption ---------------------------------------------------------
    try:
        cap = styles["Caption"]
    except KeyError:
        cap = styles.add_style("Caption", 1)
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = _rgb(WARM_GREY)
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- Metadata label (cover info table) -------------------------------
    if "Meta Label" not in [s.name for s in styles]:
        meta = styles.add_style("Meta Label", 1)
        meta.font.name = "Calibri"
        meta.font.size = Pt(9)
        meta.font.bold = True
        meta.font.color.rgb = _rgb(WARM_GREY)

    # ---- Block quote / callout -------------------------------------------
    if "Callout" not in [s.name for s in styles]:
        callout = styles.add_style("Callout", 1)
        callout.font.name = "Calibri"
        callout.font.size = Pt(10)
        callout.font.italic = True
        callout.font.color.rgb = _rgb(CHARCOAL_2)
        callout.paragraph_format.left_indent = Cm(0.8)
        callout.paragraph_format.space_before = Pt(6)
        callout.paragraph_format.space_after = Pt(6)


def _config_heading(styles, name: str, size: float, bold: bool, colour: str,
                    space_before: Pt, space_after: Pt) -> None:
    try:
        st = styles[name]
    except KeyError:
        st = styles.add_style(name, 1)
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = _rgb(colour)
    st.paragraph_format.space_before = space_before
    st.paragraph_format.space_after = space_after
    st.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover_page(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    document_code: str = "",
    study_name: str = "",
    issue: str = "1.0",
    date: str = "",
    classification: str = "Internal",
    cohort: str = "SpaceCDF",
    publisher: str = "Faculty of Engineering · School of Engineering Design and Teaching Innovation (SEDTI)",
) -> None:
    """Render a cover page in the Facilitator's Book style.

    Layout (top→bottom):
        crimson banner with uOttawa wordmark on white
        large garnet slab title
        italic serif subtitle
        cohort + year
        crimson divider rule
        publisher line, doc code · issue · date, classification
    """
    # ---- Crimson banner with uOttawa wordmark ---------------------------
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    banner.columns[0].width = Cm(17.0)
    cell = banner.rows[0].cells[0]
    _shade(cell, GARNET)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.height = Cm(3.2)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("u")
    r.font.size = Pt(36); r.font.bold = False
    r.font.color.rgb = _rgb(WHITE); r.font.name = "Calibri"
    r2 = p.add_run("Ottawa")
    r2.font.size = Pt(36); r2.font.bold = True
    r2.font.color.rgb = _rgb(WHITE); r2.font.name = "Calibri"

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(14)
    r3 = p2.add_run("UNIVERSITÉ D'OTTAWA  ·  UNIVERSITY OF OTTAWA")
    r3.font.size = Pt(9); r3.font.bold = True; r3.font.color.rgb = _rgb(WHITE)
    r3.font.name = "Calibri"

    # spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(40)

    # ---- Title block -----------------------------------------------------
    tp = doc.add_paragraph(style="Title")
    tp.add_run(title)

    # ---- Subtitle --------------------------------------------------------
    if subtitle:
        sp2 = doc.add_paragraph(style="Subtitle")
        sp2.paragraph_format.space_before = Pt(8)
        sp2.paragraph_format.space_after = Pt(20)
        sp2.add_run(subtitle)

    # ---- Cohort + year ---------------------------------------------------
    if cohort:
        cp = doc.add_paragraph()
        r4 = cp.add_run(cohort)
        r4.font.bold = True; r4.font.size = Pt(11); r4.font.color.rgb = _rgb(CHARCOAL)
    if date:
        dp = doc.add_paragraph()
        r5 = dp.add_run(date.split("-")[0] if "-" in date else date)
        r5.font.size = Pt(11); r5.font.color.rgb = _rgb(CHARCOAL)

    # ---- Spacer + crimson divider ----------------------------------------
    for _ in range(6):
        doc.add_paragraph()

    rule_p = doc.add_paragraph()
    _bottom_border(rule_p, colour=GARNET, sz=18)

    # ---- Publisher / doc code / classification --------------------------
    pp = doc.add_paragraph()
    r6 = pp.add_run(publisher)
    r6.font.bold = True; r6.font.size = Pt(10.5); r6.font.color.rgb = _rgb(GARNET)
    pp.paragraph_format.space_after = Pt(2)

    meta_bits = []
    if document_code: meta_bits.append(document_code)
    if study_name: meta_bits.append(study_name)
    meta_bits.append(f"Issue {issue}")
    if date: meta_bits.append(date)
    mp = doc.add_paragraph()
    rm = mp.add_run(" · ".join(meta_bits))
    rm.font.size = Pt(9.5); rm.font.color.rgb = _rgb(CHARCOAL)
    mp.paragraph_format.space_after = Pt(2)

    cl = doc.add_paragraph()
    rcl = cl.add_run(f"Classification: {classification}")
    rcl.font.size = Pt(9); rcl.font.italic = True; rcl.font.color.rgb = _rgb(WARM_GREY)

    # ---- Hard page break -------------------------------------------------
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document info table & change record
# ---------------------------------------------------------------------------

def add_doc_info_table(
    doc: Document,
    *,
    document_code: str,
    title: str,
    study_name: str,
    issue: str,
    date: str,
    prepared_by: str = "SpaceCDF AI Concurrent Design Facility",
    reviewed_by: str = "[TBD]",
    approved_by: str = "[TBD]",
    classification: str = "Internal",
    applies_to: str = "",
) -> None:
    """Document info matrix matching ECSS DRD style."""
    doc.add_heading("Document Information", level=1)
    rows = [
        ("Document code", document_code),
        ("Title", title),
        ("Study / Mission", study_name),
        ("Issue", issue),
        ("Date", date),
        ("Prepared by", prepared_by),
        ("Reviewed by", reviewed_by),
        ("Approved by", approved_by),
        ("Classification", classification),
    ]
    if applies_to:
        rows.append(("Applies to", applies_to))

    widths = (5.0, 12.0)  # sums to 17.0 cm — fits 17.4 cm usable width
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    _fixed_table_layout(t, widths)
    for i, (k, v) in enumerate(rows):
        c0 = t.rows[i].cells[0]; c1 = t.rows[i].cells[1]
        c0.width = Cm(widths[0]); c1.width = Cm(widths[1])
        _shade(c0, POLAR)
        _set_cell_text(c0, k, bold=True, colour=CHARCOAL, size=10)
        _set_cell_text(c1, str(v), size=10)


def add_change_record(doc: Document, entries: list[dict[str, Any]] | None = None) -> None:
    """Issue / date / summary table."""
    doc.add_heading("Change Record", level=1)
    entries = entries or [{"issue": "1.0", "date": "", "by": "SpaceCDF", "summary": "Initial issue."}]
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    _header_row(t, ["Issue", "Date", "By", "Summary of changes"])
    for e in entries:
        row = t.add_row().cells
        row[0].text = str(e.get("issue", ""))
        row[1].text = str(e.get("date", ""))
        row[2].text = str(e.get("by", ""))
        row[3].text = str(e.get("summary", ""))


# ---------------------------------------------------------------------------
# Acronyms & references
# ---------------------------------------------------------------------------

def add_acronyms_table(doc: Document, acronyms: dict[str, str]) -> None:
    doc.add_heading("Acronyms & Abbreviations", level=1)
    if not acronyms:
        doc.add_paragraph("No acronyms defined.")
        return
    widths = (3.0, 14.0)  # 17.0 cm total
    t = doc.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    _fixed_table_layout(t, widths)
    _header_row(t, ["Acronym", "Definition"])
    for cell, w in zip(t.rows[0].cells, widths):
        cell.width = Cm(w)
    for k in sorted(acronyms):
        row = t.add_row().cells
        row[0].text = k; row[0].width = Cm(widths[0])
        row[1].text = acronyms[k]; row[1].width = Cm(widths[1])


AIG_TEXT = (
    "This document was produced with the assistance of generative AI as part "
    "of the SpaceCDF Concurrent Design Facility workflow.  The design-loop "
    "convergence, agent rationales, embedded figures and document rendering "
    "are generated by the SpaceCDF backend (Python · matplotlib · python-docx) "
    "guided by ECSS, NASA SEH and SMAD-4 references.  Editorial framing, "
    "worked examples and pedagogical commentary remain owned by the SpaceCDF "
    "teaching team."
)

AIG_ATTRIBUTION = (
    "Attribution follows the AIG (Assisted by Generative AI) framework — "
    "Peters (2023), Logos IA-EN, CC BY-NC-SA 4.0 — "
    "https://mpeters.uqo.ca/en/logos-ia-en-peters-2023/"
)

AIG_DOWNSTREAM = (
    "Any course deliverable that incorporates content from this document, or "
    "from any document exported by SpaceCDF, must in turn carry the AIG "
    "badge and a short note describing how generative AI was used."
)


def add_aig_acknowledgement(doc: Document, *, heading: str = "Acknowledgement — Generative AI (AIG)") -> None:
    """Insert the standardised AIG / Peters (2023) attribution block.

    The block is short (three paragraphs) and idempotent — call it once per
    document, typically immediately after the change record or in the front
    matter for non-DID documents.
    """
    doc.add_heading(heading, level=1)
    p = doc.add_paragraph(AIG_TEXT)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    r = p.add_run(AIG_ATTRIBUTION)
    r.font.italic = True
    p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph(AIG_DOWNSTREAM)


def add_reference_list(doc: Document, references: list[dict[str, str]], heading: str = "Reference Documents") -> None:
    doc.add_heading(heading, level=1)
    if not references:
        doc.add_paragraph("None.")
        return
    widths = (1.8, 4.0, 11.2)  # 17.0 cm total
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.autofit = False
    _fixed_table_layout(t, widths)
    _header_row(t, ["ID", "Reference", "Title"])
    for cell, w in zip(t.rows[0].cells, widths):
        cell.width = Cm(w)
    for i, ref in enumerate(references, 1):
        row = t.add_row().cells
        row[0].text = ref.get("id", f"RD-{i:02d}")
        row[1].text = ref.get("ref", "")
        row[2].text = ref.get("title", "")
        for cell, w in zip(row, widths):
            cell.width = Cm(w)


# ---------------------------------------------------------------------------
# Table of contents (Word field)
# ---------------------------------------------------------------------------

def add_toc(doc: Document, levels: str = "1-3") -> None:
    """Insert a Word native ToC field; updates on open."""
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select 'Update Field' to populate."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Page furniture (header rule, bilingual footer)
# ---------------------------------------------------------------------------

def add_page_furniture(
    doc: Document,
    *,
    running_title: str = "SpaceCDF",
    document_code: str = "",
    footer_left: str = "SpaceCDF",
    footer_right: str = "uOttawa SEDTI",
) -> None:
    """Add running header (left: title, right: doc code) with crimson rule,
    plus three-column bilingual footer: footer_left | Page X of / de Y | footer_right.
    Applies to all sections in the document.
    """
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        h_tbl = _three_col_table(header.add_paragraph().part, header, widths=(Cm(8), Cm(8.7)))
        # Actually python-docx headers are paragraph-based; simpler: clear and build
        for p in list(header.paragraphs):
            p.clear()
        hp = header.paragraphs[0]
        tab_stops = hp.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)
        rL = hp.add_run(running_title)
        rL.font.bold = True; rL.font.size = Pt(9); rL.font.color.rgb = _rgb(CHARCOAL)
        hp.add_run("\t")
        if document_code:
            rR = hp.add_run(document_code)
            rR.font.size = Pt(9); rR.font.color.rgb = _rgb(WARM_GREY)
        _bottom_border(hp, colour=GARNET, sz=12)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p.clear()
        fp = footer.paragraphs[0]
        tab_stops = fp.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(8.5), WD_ALIGN_PARAGRAPH.CENTER)
        tab_stops.add_tab_stop(Cm(16.0), WD_ALIGN_PARAGRAPH.RIGHT)

        rfl = fp.add_run(footer_left)
        rfl.font.size = Pt(8); rfl.font.color.rgb = _rgb(CHARCOAL)
        fp.add_run("\t")
        rfm1 = fp.add_run("Page ")
        rfm1.font.size = Pt(8); rfm1.font.color.rgb = _rgb(WARM_GREY)
        _insert_field(fp, "PAGE")
        rfm2 = fp.add_run(" of / de ")
        rfm2.font.size = Pt(8); rfm2.font.color.rgb = _rgb(WARM_GREY)
        _insert_field(fp, "NUMPAGES")
        fp.add_run("\t")
        rfr = fp.add_run(footer_right)
        rfr.font.size = Pt(8); rfr.font.bold = True; rfr.font.color.rgb = _rgb(GARNET)


def _insert_field(paragraph, instr: str) -> None:
    """Insert a Word field (PAGE / NUMPAGES) into a paragraph as small grey text."""
    run = paragraph.add_run()
    run.font.size = Pt(8); run.font.color.rgb = _rgb(WARM_GREY)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    fld = OxmlElement("w:instrText"); fld.set(qn("xml:space"), "preserve"); fld.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(fld); run._r.append(sep); run._r.append(txt); run._r.append(end)


def _three_col_table(part, container, widths):
    # Placeholder retained for future use; we currently use tab stops.
    return None


# ---------------------------------------------------------------------------
# Bookmarked headings + numbered headings
# ---------------------------------------------------------------------------

_BOOKMARK_ID = [1000]


def bookmarked_heading(doc: Document, text: str, *, level: int = 1, bookmark: str | None = None):
    """Add a heading and (optionally) wrap it in a bookmark for cross-references."""
    h = doc.add_heading(text, level=level)
    if bookmark:
        _BOOKMARK_ID[0] += 1
        bid = str(_BOOKMARK_ID[0])
        start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), bid); start.set(qn("w:name"), bookmark)
        end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), bid)
        h._p.insert(0, start)
        h._p.append(end)
    return h


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

def styled_table(
    doc: Document,
    *,
    headers: list[str],
    rows: list[list[Any]],
    col_widths_cm: list[float] | None = None,
    header_fill: str = GARNET,
    header_font_colour: str = WHITE,
    alt_fill: str = POLAR,
    status_col_index: int | None = None,
) -> Table:
    """Build a SpaceCDF-styled table with crimson header & alternating rows.

    If ``status_col_index`` is provided, cells in that column are shaded by
    their value mapped through ``STATUS_COLOURS``.
    """
    n_cols = len(headers)
    t = doc.add_table(rows=1, cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    if col_widths_cm:
        # Clamp the requested widths to the usable page width so a too-wide
        # column spec is gracefully scaled instead of overflowing the page.
        col_widths_cm = list(col_widths_cm[:n_cols])
        total = sum(col_widths_cm)
        if total > USABLE_WIDTH_CM:
            scale = USABLE_WIDTH_CM / total
            col_widths_cm = [w * scale for w in col_widths_cm]
        t.autofit = False
        _fixed_table_layout(t, col_widths_cm)
        # Set widths on the header row cells too — Word otherwise lets them
        # stretch to content
        for ci, w in enumerate(col_widths_cm):
            t.rows[0].cells[ci].width = Cm(w)
    else:
        t.autofit = True

    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, header_fill)
        _set_cell_text(c, h, bold=True, colour=header_font_colour, size=10)

    # Body
    for ri, row in enumerate(rows):
        row_cells = t.add_row().cells
        for ci, val in enumerate(row[:n_cols]):
            cell = row_cells[ci]
            txt = "" if val is None else (f"{val:.2f}" if isinstance(val, float) else str(val))
            _set_cell_text(cell, txt, size=10)
            if col_widths_cm:
                cell.width = Cm(col_widths_cm[ci])
            if ri % 2 == 1:
                _shade(cell, alt_fill)
            if status_col_index is not None and ci == status_col_index:
                key = str(val).lower().replace(" ", "_") if val is not None else ""
                colour = STATUS_COLOURS.get(key)
                if colour:
                    _shade(cell, colour)
    return t


def _header_row(t: Table, headers: list[str]) -> None:
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shade(c, GARNET)
        _set_cell_text(c, h, bold=True, colour=WHITE, size=10)


# ---------------------------------------------------------------------------
# Figures & callouts
# ---------------------------------------------------------------------------

_FIG_COUNTER = [0]
_TBL_COUNTER = [0]


def reset_counters() -> None:
    _FIG_COUNTER[0] = 0
    _TBL_COUNTER[0] = 0


def add_figure(doc: Document, png_bytes: bytes, *, caption: str, width_cm: float = 14.0) -> None:
    """Insert a centred figure with a "Figure N — caption" line beneath."""
    _FIG_COUNTER[0] += 1
    from io import BytesIO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(BytesIO(png_bytes), width=Cm(width_cm))
    cap = doc.add_paragraph(style="Caption")
    r1 = cap.add_run(f"Figure {_FIG_COUNTER[0]} — ")
    r1.font.bold = True; r1.font.italic = True; r1.font.color.rgb = _rgb(GARNET)
    r2 = cap.add_run(caption)
    r2.font.italic = True; r2.font.color.rgb = _rgb(WARM_GREY)


def add_table_caption(doc: Document, caption: str) -> None:
    _TBL_COUNTER[0] += 1
    p = doc.add_paragraph(style="Caption")
    r1 = p.add_run(f"Table {_TBL_COUNTER[0]} — ")
    r1.font.bold = True; r1.font.italic = True; r1.font.color.rgb = _rgb(GARNET)
    r2 = p.add_run(caption)
    r2.font.italic = True; r2.font.color.rgb = _rgb(WARM_GREY)


def add_callout(doc: Document, text: str, *, kind: str = "note") -> None:
    """A soft-shaded callout box (note / warning / info)."""
    fill = {"note": POLAR, "warning": "FFEB9C", "info": "E7F0F7"}.get(kind, POLAR)
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.columns[0].width = Cm(17.0)
    c = t.rows[0].cells[0]
    _shade(c, fill)
    _set_cell_text(c, text, size=10, italic=True, colour=CHARCOAL_2)
    # left border accent
    _left_border(c, colour=GARNET, sz=24)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Low-level OXML helpers
# ---------------------------------------------------------------------------

def _shade(cell: _Cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False, italic: bool = False,
                   colour: str = CHARCOAL, size: float = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = _rgb(colour); r.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _bottom_border(paragraph, *, colour: str, sz: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pBdr = p_pr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); p_pr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), colour)
    pBdr.append(bottom)


def _fixed_table_layout(table: Table, col_widths_cm) -> None:
    """Force the table to use a fixed (non-autofit) layout so cell widths stick.

    Word's default behaviour is to stretch cells around their content.  For
    documents that need to fit a fixed page width, fixed-layout tables are
    much more predictable.
    """
    tblPr = table._tbl.tblPr
    # Remove any existing layout element to avoid duplicates
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # Also write the explicit table width (sum of columns) so Word doesn't
    # rebalance the table across all of "auto".
    total = sum(col_widths_cm)
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(total * 567)))  # 1 cm = 567 twips
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    for i, w in enumerate(col_widths_cm):
        if i < len(table.columns):
            table.columns[i].width = Cm(w)


def _left_border(cell: _Cell, *, colour: str, sz: int = 24) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = tc_pr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders"); tc_pr.append(tcBorders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4"); left.set(qn("w:color"), colour)
    tcBorders.append(left)
