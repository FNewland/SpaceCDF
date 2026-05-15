"""SpaceCDF -- Word Document Generator.

Generates editable .docx files for all ECSS documents using python-docx.
Each document is populated from the live design state.

Two rendering paths:
  1. Legacy per-type generators (generate_mrd_docx, etc.) — hardcoded sections.
  2. Universal DID renderer (generate_did_docx) — takes any DID JSON with
     {sections: [{number, title, subsections: [{number, title, content}]}]}
     and renders rich content (tables, bullets, mixed text).

Supports: MRD, TS, VP, ConOps, SEMP, RMP, IRD, Test Plan, BOM.

Visual style: uOttawa SpaceCDF Facilitator's Book.  All cover pages,
headers, footers, ToCs and tables route through
``spacecdf_agents.exporters.docs.theme`` so every document shares the
garnet course identity.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# uOttawa SpaceCDF course-style theme (shared with the design-review generator).
from spacecdf_agents.exporters.docs import theme as _theme


# ---------------------------------------------------------------------------
# Colour palette — uOttawa Horizon (kept as RGBColor tuples for the legacy
# helpers; identical to the theme module's constants).
# ---------------------------------------------------------------------------

_CLR_HEADER_BG = RGBColor(0x8f, 0x00, 0x1a)   # uOttawa garnet — table headers
_CLR_ALT_ROW = RGBColor(0xf2, 0xf2, 0xf2)     # polar grey — alternate rows
_CLR_WHITE = RGBColor(255, 255, 255)
_CLR_TITLE = RGBColor(0x8f, 0x00, 0x1a)       # garnet — headings & titles
_CLR_SUBTITLE = RGBColor(0x80, 0x74, 0x6c)    # warm grey — captions & metadata
_CLR_BODY = RGBColor(0x2d, 0x2d, 0x2c)        # charcoal — body text
_CLR_ACCENT = RGBColor(0x67, 0x79, 0x6c)      # uOttawa green — accent


# ---------------------------------------------------------------------------
# Document setup helpers
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document):
    """Apply the uOttawa SpaceCDF Facilitator's Book styles to the document."""
    _theme.apply_styles(doc)


def _setup_margins(doc: Document):
    """Set the SpaceCDF page geometry on every section."""
    _theme._set_page_geometry(doc)


def _add_footer(doc: Document, study_name: str, document_code: str = ""):
    """Bilingual EN/FR running header + footer matching the course style."""
    _theme.add_page_furniture(
        doc,
        running_title=f"SpaceCDF \u2014 {study_name}",
        document_code=document_code,
        footer_left=f"SpaceCDF \u00b7 {study_name}",
        footer_right="uOttawa SEDTI",
    )


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover_page(
    doc: Document,
    title: str,
    standard: str,
    study_name: str,
    date_str: str | None = None,
    status: str = "DRAFT",
):
    """uOttawa SpaceCDF-style cover page (delegates to the shared theme)."""
    date_str = date_str or datetime.now().strftime("%Y-%m-%d")
    _theme.add_cover_page(
        doc,
        title=title,
        subtitle=standard,
        document_code=standard,
        study_name=study_name,
        issue="1.0",
        date=date_str,
        classification=status if status != "DRAFT" else "Internal · DRAFT",
        cohort="SpaceCDF",
        publisher=("Faculty of Engineering · School of Engineering Design "
                   "and Teaching Innovation (SEDTI)"),
    )


# ---------------------------------------------------------------------------
# Table of Contents placeholder
# ---------------------------------------------------------------------------

def _add_toc_placeholder(doc: Document):
    """Insert a Word-native ToC field (delegates to the shared theme)."""
    _theme.add_toc(doc)


# ---------------------------------------------------------------------------
# Content detection & rendering
# ---------------------------------------------------------------------------

_PIPE_TABLE_RE = re.compile(r'^\|(.+)\|$')
_SEPARATOR_RE = re.compile(r'^\|\s*[-:]+[\s|:-]*\|$')


def _is_pipe_table_line(line: str) -> bool:
    """Check if a line is a pipe-delimited table row."""
    return bool(_PIPE_TABLE_RE.match(line.strip()))


def _is_separator_line(line: str) -> bool:
    """Check if a line is a markdown table separator (|---|---|)."""
    return bool(_SEPARATOR_RE.match(line.strip()))


def _is_bullet_line(line: str) -> bool:
    """Check if a line starts with - or * (bullet list)."""
    stripped = line.strip()
    return stripped.startswith("- ") or stripped.startswith("* ")


def _parse_pipe_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse pipe-delimited table lines into headers and rows.

    Returns (headers, rows) where each row is a list of cell strings.
    Skips separator lines.
    """
    headers: list[str] = []
    rows: list[list[str]] = []
    header_done = False

    for line in lines:
        stripped = line.strip()
        if _is_separator_line(stripped):
            header_done = True
            continue
        match = _PIPE_TABLE_RE.match(stripped)
        if not match:
            continue
        cells = [c.strip() for c in match.group(1).split('|')]
        if not header_done:
            headers = cells
            header_done = True
        else:
            rows.append(cells)

    return headers, rows


def _set_cell_shading(cell, color: RGBColor):
    """Set background shading on a table cell."""
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _render_formatted_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Render a table with styled header and alternating row colours."""
    if not headers:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        run.font.color.rgb = _CLR_WHITE
        _set_cell_shading(cell, _CLR_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(num_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            # Alternating row colour
            if r_idx % 2 == 1:
                _set_cell_shading(cell, _CLR_ALT_ROW)

    doc.add_paragraph()  # spacing after table


def _render_bullet(doc: Document, text: str):
    """Add a bullet-point paragraph."""
    # Strip leading "- " or "* "
    clean = text.strip()
    if clean.startswith("- "):
        clean = clean[2:]
    elif clean.startswith("* "):
        clean = clean[2:]

    p = doc.add_paragraph(style='List Bullet')
    # Handle bold segments with **text** markers
    _add_formatted_run(p, clean)


def _add_formatted_run(paragraph, text: str):
    """Add text to a paragraph, handling **bold** markdown markers."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'


def _render_rich_content(doc: Document, content: str):
    """Parse and render mixed content: text, pipe tables, and bullet lists.

    Handles content that contains a mix of plain paragraphs, pipe-delimited
    tables, and bullet lines, rendering each block with the appropriate style.
    """
    if not content or not content.strip():
        return

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Accumulate consecutive pipe-table lines
        if _is_pipe_table_line(line):
            table_lines: list[str] = []
            while i < len(lines) and (_is_pipe_table_line(lines[i]) or _is_separator_line(lines[i])):
                table_lines.append(lines[i])
                i += 1
            headers, rows = _parse_pipe_table(table_lines)
            if headers:
                _render_formatted_table(doc, headers, rows)
            continue

        # Skip standalone separator lines (edge case)
        if _is_separator_line(line):
            i += 1
            continue

        # Bullet list items
        if _is_bullet_line(line):
            _render_bullet(doc, line)
            i += 1
            continue

        # Plain text paragraph (skip empty lines)
        stripped = line.strip()
        if stripped:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)
        i += 1


# ---------------------------------------------------------------------------
# Heading level helper
# ---------------------------------------------------------------------------

def _heading_level(number: str) -> int:
    """Determine heading level from section number (1 -> 1, 1.1 -> 2, 1.1.1 -> 3)."""
    dots = number.count('.')
    if dots == 0:
        return 1
    elif dots == 1:
        return 2
    else:
        return 3


# ---------------------------------------------------------------------------
# Legacy helpers (preserved for backward compatibility)
# ---------------------------------------------------------------------------

def _add_title_page(doc: Document, title: str, subtitle: str, study_name: str):
    """Add a formatted title page (legacy — now wraps _add_cover_page)."""
    _add_cover_page(doc, title, subtitle, study_name)


def _add_section(doc: Document, number: str, title: str, content: str = ""):
    """Add a numbered section heading with content (legacy)."""
    level = _heading_level(number)
    doc.add_heading(f"{number} {title}", level=level)
    if content:
        _render_rich_content(doc, content)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a formatted table (legacy — now wraps _render_formatted_table)."""
    _render_formatted_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Universal DID-to-DOCX renderer
# ---------------------------------------------------------------------------

def generate_did_docx(did_json: dict[str, Any]) -> bytes:
    """Convert a DID JSON response into a styled .docx file.

    Works for ALL document types (MRD, TS, IRD, ConOps, SEMP, RMP, Test Plan,
    etc.) since they all share the structure:
      {
        document: str,
        standard: str,
        study_name: str,
        document_code: str            (optional — defaults to standard)
        issue: str                    (optional — defaults to "1.0")
        classification: str           (optional — defaults to "Internal")
        acronyms: dict[str,str]       (optional — adds an acronym table)
        applicable_documents: list    (optional — list of {id,ref,title})
        reference_documents: list     (optional — list of {id,ref,title})
        change_record: list           (optional — list of {issue,date,by,summary})
        sections: [
          {number, title, content, subsections: [...]}
        ]
      }

    Content strings may contain:
      - Pipe-delimited tables
      - Bullet lists (lines starting with "- " or "* ")
      - Bold markers (**text**)
      - Plain text paragraphs
      - Any mix of the above
    """
    doc = _theme.new_document()
    _theme.reset_counters()

    title = did_json.get("document", "SpaceCDF Document")
    standard = did_json.get("standard", "")
    study_name = did_json.get("study_name", "Unnamed Mission")
    generated = did_json.get("generated", "")
    date_str = generated[:10] if generated else datetime.now().strftime("%Y-%m-%d")
    document_code = did_json.get("document_code") or standard or "SCDF-DOC"
    issue = did_json.get("issue", "1.0")
    classification = did_json.get("classification", "Internal · DRAFT")

    # Cover page
    _theme.add_cover_page(
        doc,
        title=title,
        subtitle=standard,
        document_code=document_code,
        study_name=study_name,
        issue=issue,
        date=date_str,
        classification=classification,
        cohort="SpaceCDF",
        publisher=("Faculty of Engineering · School of Engineering Design "
                   "and Teaching Innovation (SEDTI)"),
    )

    # Page furniture (header rule + bilingual footer with Page X of / de Y)
    _theme.add_page_furniture(
        doc,
        running_title=f"SpaceCDF — {title}",
        document_code=document_code,
        footer_left=f"{document_code} · {study_name}",
        footer_right="uOttawa SEDTI",
    )

    # Document information & change record
    _theme.add_doc_info_table(
        doc, document_code=document_code, title=title,
        study_name=study_name, issue=issue, date=date_str,
        classification=classification, applies_to=study_name,
    )
    if did_json.get("change_record"):
        _theme.add_change_record(doc, did_json["change_record"])
    else:
        _theme.add_change_record(doc, [
            {"issue": issue, "date": date_str, "by": "SpaceCDF",
             "summary": f"Initial issue of {title}."}
        ])

    # AIG (Peters 2023) acknowledgement — applies to every exported document
    _theme.add_aig_acknowledgement(doc)

    # Acronyms & references (optional)
    if did_json.get("acronyms"):
        _theme.add_acronyms_table(doc, did_json["acronyms"])
    if did_json.get("applicable_documents"):
        _theme.add_reference_list(doc, did_json["applicable_documents"],
                                  heading="Applicable Documents")
    if did_json.get("reference_documents"):
        _theme.add_reference_list(doc, did_json["reference_documents"],
                                  heading="Reference Documents")

    # Table of contents
    _theme.add_toc(doc)

    # Render sections
    sections = did_json.get("sections", [])
    for section in sections:
        sec_number = section.get("number", "")
        sec_title = section.get("title", "")
        sec_level = _heading_level(sec_number)

        doc.add_heading(f"{sec_number} {sec_title}", level=sec_level)

        # Some sections may have direct content (top-level)
        sec_content = section.get("content", "")
        if sec_content:
            _render_rich_content(doc, sec_content)

        # Subsections
        subsections = section.get("subsections", [])
        for sub in subsections:
            sub_number = sub.get("number", "")
            sub_title = sub.get("title", "")
            sub_level = _heading_level(sub_number)
            sub_content = sub.get("content", "")

            doc.add_heading(f"{sub_number} {sub_title}", level=sub_level)
            if sub_content:
                _render_rich_content(doc, sub_content)

            # Nested subsections (if any)
            nested = sub.get("subsections", [])
            for nested_sub in nested:
                n_number = nested_sub.get("number", "")
                n_title = nested_sub.get("title", "")
                n_content = nested_sub.get("content", "")
                n_level = _heading_level(n_number)

                doc.add_heading(f"{n_number} {n_title}", level=n_level)
                if n_content:
                    _render_rich_content(doc, n_content)

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# BOM-specific DOCX renderer
# ---------------------------------------------------------------------------

def _render_bom_docx(bom_data: dict[str, Any]) -> bytes:
    """Render a Bill of Materials JSON as a styled Word document.

    Expected bom_data structure (from bom_generator):
      {
        title: str,
        generated: str,
        form_factor: str,
        groups: {subsystem_name: [line_items]},
        lines: [all_line_items],
        summary: {total_lines, total_mass_kg, total_power_w, total_cost_keur, ...},
        procurement_notes: [str],
        segment_totals: {segment: {mass_kg, power_w, cost_keur, items}},
      }
    """
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    title = bom_data.get("title", "Bill of Materials")
    study_name = title.replace("Bill of Materials \u2014 ", "").replace("SpaceCDF Bill of Materials \u2014 ", "")
    generated = bom_data.get("generated", "")
    date_str = generated[:10] if generated else datetime.now().strftime("%Y-%m-%d")

    # Cover page
    _add_cover_page(doc, "Bill of Materials", "SpaceCDF Procurement Package", study_name, date_str=date_str, status="DRAFT")
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    summary = bom_data.get("summary", {})
    groups = bom_data.get("groups", {})
    segment_totals = bom_data.get("segment_totals", {})
    procurement_notes = bom_data.get("procurement_notes", [])

    # --- Section 1: Summary ---
    doc.add_heading("1 Summary", level=1)

    summary_headers = ["Metric", "Value"]
    summary_rows = [
        ["Total line items", str(summary.get("total_lines", 0))],
        ["Total dry mass", f"{summary.get('total_mass_kg', 0):.3f} kg"],
        ["Total power", f"{summary.get('total_power_w', 0):.1f} W"],
        ["Total cost", f"{summary.get('total_cost_keur', 0):.1f} kEUR ({summary.get('total_cost_eur', 0):.0f} EUR)"],
        ["Mean TRL", str(summary.get("mean_trl", "N/A"))],
        ["BOM completeness", f"{summary.get('completeness_percent', 0):.0f}%"],
        ["Critical path", f"{summary.get('critical_path_weeks', 0)} weeks"],
        ["ITAR/EAR items", str(summary.get("itar_items", 0))],
        ["Low-TRL items (< 7)", str(summary.get("low_trl_items", 0))],
        ["Form factor", bom_data.get("form_factor", "N/A")],
    ]
    _render_formatted_table(doc, summary_headers, summary_rows)

    # --- Section 2: Segment Totals ---
    if segment_totals:
        doc.add_heading("2 Segment Totals", level=1)
        seg_headers = ["Segment", "Items", "Mass (kg)", "Power (W)", "Cost (kEUR)"]
        seg_rows = []
        for seg_name, seg_data in segment_totals.items():
            seg_rows.append([
                seg_name.title(),
                str(seg_data.get("items", 0)),
                f"{seg_data.get('mass_kg', 0):.3f}",
                f"{seg_data.get('power_w', 0):.1f}",
                f"{seg_data.get('cost_keur', 0):.1f}",
            ])
        _render_formatted_table(doc, seg_headers, seg_rows)

    # --- Section 3: Grouped BOM ---
    doc.add_heading("3 Bill of Materials by Subsystem", level=1)

    group_idx = 0
    for subsys_name, lines in groups.items():
        group_idx += 1
        doc.add_heading(f"3.{group_idx} {subsys_name}", level=2)

        bom_headers = [
            "Line", "Name", "Qty", "Mass (kg)",
            "Power (W)", "Cost (kEUR)", "TRL", "Model",
            "Procurement", "Lead (wk)",
        ]
        bom_rows = []
        for line in lines:
            bom_rows.append([
                str(line.get("line", "")),
                line.get("name", ""),
                str(line.get("quantity", 1)),
                f"{line.get('total_mass_kg', 0):.3f}",
                f"{line.get('total_power_w', 0):.1f}",
                f"{line.get('total_cost_keur', 0):.1f}",
                str(line.get("trl", "")),
                line.get("model_level", ""),
                line.get("procurement_status", ""),
                str(line.get("lead_time_weeks", "")),
            ])

        _render_formatted_table(doc, bom_headers, bom_rows)

        # Subsystem sub-totals
        sub_mass = sum(l.get("total_mass_kg", 0) for l in lines)
        sub_power = sum(l.get("total_power_w", 0) for l in lines)
        sub_cost = sum(l.get("total_cost_keur", 0) for l in lines)
        p = doc.add_paragraph()
        run = p.add_run(
            f"Subtotal: {len(lines)} items, "
            f"{sub_mass:.3f} kg, {sub_power:.1f} W, {sub_cost:.1f} kEUR"
        )
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = 'Calibri'

    # --- Section 4: Procurement Notes ---
    doc.add_heading("4 Procurement Notes", level=1)
    if procurement_notes:
        for note in procurement_notes:
            _render_bullet(doc, f"- {note}")
    else:
        doc.add_paragraph("No procurement notes.")

    # --- Section 5: Export Control ---
    doc.add_heading("5 Export Control Summary", level=1)
    export_items = [l for l in bom_data.get("lines", []) if l.get("export_control") not in ("none", "", None)]
    if export_items:
        exp_headers = ["Line", "Name", "Manufacturer", "Export Control"]
        exp_rows = [
            [
                str(l.get("line", "")),
                l.get("name", ""),
                l.get("manufacturer", ""),
                l.get("export_control", ""),
            ]
            for l in export_items
        ]
        _render_formatted_table(doc, exp_headers, exp_rows)
    else:
        doc.add_paragraph("No items with export control restrictions identified.")

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Legacy per-type generators (preserved for backward compat)
# ---------------------------------------------------------------------------

def generate_mrd_docx(
    study_name: str = "",
    mission_need: dict[str, Any] | None = None,
    requirements: list[dict[str, Any]] | None = None,
    orbit: dict[str, Any] | None = None,
    elements: list[dict[str, Any]] | None = None,
) -> bytes:
    """Generate Mission Requirements Document as .docx bytes."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)
    mn = mission_need or {}
    reqs = requirements or []
    orb = orbit or {}
    elems = elements or []

    _add_cover_page(doc, "Mission Requirements Document", "ECSS-E-ST-10C Annex A", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    # ── Section 1: Executive Summary ──
    _add_section(doc, "1", "Executive Summary")
    problem = mn.get("problem_statement", "")
    sel_rationale = mn.get("selection_rationale", "")
    if problem:
        summary_text = problem
        if sel_rationale:
            summary_text += f"\n\n{sel_rationale}"
        _render_rich_content(doc, summary_text)
    else:
        doc.add_paragraph(
            f"This Mission Requirements Document (MRD) defines the mission-level and "
            f"system-level requirements for the {study_name} mission. It captures "
            f"stakeholder needs, mission objectives, the alternatives analysis, and "
            f"the traceable requirements baseline."
        )

    # ── Section 2: Mission Need ──
    _add_section(doc, "2", "Mission Need")
    _add_section(doc, "2.1", "Problem Statement",
                 mn.get("problem_statement", "[To be defined]"))
    _add_section(doc, "2.2", "Operational Context",
                 mn.get("operational_context", "[To be defined]"))

    # ── Section 3: Stakeholder Register ──
    _add_section(doc, "3", "Stakeholder Register")
    stakeholders = mn.get("stakeholders", [])
    if stakeholders:
        _add_table(doc,
                   ["Name", "Role", "Key Needs", "Priority"],
                   [
                       [
                           s.get("name", ""),
                           s.get("role", ""),
                           ", ".join(s.get("needs", [])) if isinstance(s.get("needs"), list) else str(s.get("needs", "")),
                           s.get("priority", ""),
                       ]
                       for s in stakeholders
                   ])
    else:
        doc.add_paragraph("[Stakeholder register to be populated]")

    # ── Section 4: Mission Objectives ──
    _add_section(doc, "4", "Mission Objectives")
    objectives = mn.get("objectives", [])
    if objectives:
        _add_table(doc,
                   ["ID", "Objective", "Priority", "Measurable Criterion", "Status"],
                   [
                       [
                           o.get("id", ""),
                           o.get("text", ""),
                           o.get("priority", ""),
                           o.get("measurable_criterion", ""),
                           o.get("status", "proposed"),
                       ]
                       for o in objectives
                   ])
    else:
        doc.add_paragraph("[Objectives to be defined]")

    # ── Section 5: Alternatives Analysis ──
    _add_section(doc, "5", "Alternatives Analysis")
    alternatives = mn.get("alternatives", [])
    if alternatives:
        _add_table(doc,
                   ["Name", "Type", "Pros", "Cons", "Feasibility", "Decision"],
                   [
                       [
                           a.get("name", ""),
                           a.get("type", ""),
                           "; ".join(a.get("pros", [])) if isinstance(a.get("pros"), list) else str(a.get("pros", "")),
                           "; ".join(a.get("cons", [])) if isinstance(a.get("cons"), list) else str(a.get("cons", "")),
                           f"{a.get('feasibility_score', 0):.1f}" if a.get("feasibility_score") else "",
                           a.get("decision", ""),
                       ]
                       for a in alternatives
                   ])
    else:
        doc.add_paragraph("[Alternatives analysis to be performed]")

    # ── Section 6: Selected Concept ──
    _add_section(doc, "6", "Selected Concept")
    sel_alt_id = mn.get("selected_alternative_id")
    selected_alt = None
    if sel_alt_id and alternatives:
        selected_alt = next((a for a in alternatives if a.get("id") == sel_alt_id), None)
    if selected_alt:
        p = doc.add_paragraph()
        run = p.add_run(f"Selected concept: {selected_alt.get('name', '')}")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        if selected_alt.get("description"):
            doc.add_paragraph(selected_alt["description"])
    _add_section(doc, "6.1", "Selection Rationale",
                 mn.get("selection_rationale", "[To be documented]"))

    # ── Section 7: Mission Requirements ──
    _add_section(doc, "7", "Mission Requirements")
    mission_reqs = [r for r in reqs if r.get("level", "").lower() == "mission"]
    if mission_reqs:
        _add_table(doc,
                   ["ID", "Code", "Type", "Requirement Text", "Verification"],
                   [
                       [
                           r.get("id", "")[:12],
                           r.get("code", ""),
                           r.get("type", r.get("req_type", "")),
                           r.get("text", ""),
                           r.get("verification_method", ""),
                       ]
                       for r in mission_reqs[:80]
                   ])
    elif reqs:
        doc.add_paragraph(
            f"No requirements are tagged with level='mission'. "
            f"Showing all {len(reqs)} requirements below in Section 8."
        )
    else:
        doc.add_paragraph("[Requirements to be generated from objectives]")

    # ── Section 8: System Requirements ──
    _add_section(doc, "8", "System Requirements")
    system_reqs = [r for r in reqs if r.get("level", "").lower() == "system"]
    if system_reqs:
        _add_table(doc,
                   ["ID", "Code", "Type", "Requirement Text", "Verification"],
                   [
                       [
                           r.get("id", "")[:12],
                           r.get("code", ""),
                           r.get("type", r.get("req_type", "")),
                           r.get("text", ""),
                           r.get("verification_method", ""),
                       ]
                       for r in system_reqs[:80]
                   ])
    elif not mission_reqs and reqs:
        # No level-tagged reqs -- show all
        _add_table(doc,
                   ["ID", "Code", "Level", "Type", "Requirement Text", "Verification"],
                   [
                       [
                           r.get("id", "")[:12],
                           r.get("code", ""),
                           r.get("level", ""),
                           r.get("type", r.get("req_type", "")),
                           r.get("text", ""),
                           r.get("verification_method", ""),
                       ]
                       for r in reqs[:80]
                   ])
    else:
        doc.add_paragraph("[System requirements to be derived from mission requirements]")

    # ── Section 9: Orbit & Payload Summary ──
    _add_section(doc, "9", "Orbit & Payload Summary")
    if orb:
        orbit_rows = [
            ["Orbit Type", str(orb.get("orbit_type", "TBD"))],
            ["Altitude", f"{orb.get('altitude_km', 'TBD')} km"],
            ["Inclination", f"{orb.get('inclination_deg', 'TBD')} deg"],
            ["Eccentricity", str(orb.get("eccentricity", "TBD"))],
            ["Design Lifetime", f"{orb.get('design_lifetime_years', 'TBD')} years"],
        ]
        if orb.get("period_min"):
            orbit_rows.append(["Orbital Period", f"{orb['period_min']} min"])
        if orb.get("eclipse_fraction"):
            orbit_rows.append(["Eclipse Fraction", f"{orb['eclipse_fraction']}"])
        _add_table(doc, ["Parameter", "Value"], orbit_rows)
    else:
        doc.add_paragraph("[Orbit parameters to be defined]")

    # Payload elements
    payload_elems = [e for e in elems if e.get("subsystem_domain") == "payload"]
    if payload_elems:
        _add_section(doc, "9.1", "Payloads")
        _add_table(doc,
                   ["Name", "Mass (kg)", "Power (W)", "Data Rate"],
                   [
                       [
                           e.get("name", ""),
                           str(e.get("mass_kg", "")),
                           str(e.get("power_avg_w", "")),
                           str((e.get("performance") or {}).get("data_rate_mbps", "")),
                       ]
                       for e in payload_elems
                   ])

    # ── Section 10: Constraints ──
    _add_section(doc, "10", "Constraints")
    _add_section(doc, "10.1", "Programmatic Constraints",
                 "[Budget, schedule, launch date constraints]")
    _add_section(doc, "10.2", "Technical Constraints",
                 "[Orbit, mass, interfaces, regulatory]")
    _add_section(doc, "10.3", "Regulatory Constraints",
                 "Space debris mitigation per ECSS-U-AS-10C Rev.2. ITU frequency coordination required.")

    # Return as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_conops_docx(
    study_name: str = "",
    mission_need: dict[str, Any] | None = None,
    conops: dict[str, Any] | None = None,
    orbit: dict[str, Any] | None = None,
    elements: list[dict[str, Any]] | None = None,
) -> bytes:
    """Generate Concept of Operations document as .docx bytes."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)
    mn = mission_need or {}
    ops = conops or {}
    orb = orbit or {}
    elems = elements or []

    _add_cover_page(doc, "Concept of Operations", "NASA SEH Appendix S", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    # ── Section 1: Mission Overview ──
    _add_section(doc, "1", "Mission Overview")
    _add_section(doc, "1.1", "Purpose",
                 f"This document describes the concept of operations for the {study_name} "
                 f"mission, defining how the system will be operated throughout its "
                 f"lifecycle to meet mission objectives.")
    problem = mn.get("problem_statement", "")
    sel_rationale = mn.get("selection_rationale", "")
    overview_parts = []
    if problem:
        overview_parts.append(problem)
    if sel_rationale:
        overview_parts.append(f"\n\n**Selected Concept:** {sel_rationale}")
    _add_section(doc, "1.2", "Mission Description",
                 "\n".join(overview_parts) if overview_parts else "[To be defined]")
    op_context = mn.get("operational_context", "")
    if op_context:
        _add_section(doc, "1.3", "Operational Context", op_context)

    # ── Section 2: Orbital Parameters ──
    _add_section(doc, "2", "Orbital Parameters")
    if orb:
        orbit_rows = [
            ["Orbit Type", str(orb.get("orbit_type", "TBD"))],
            ["Altitude", f"{orb.get('altitude_km', 'TBD')} km"],
            ["Inclination", f"{orb.get('inclination_deg', 'TBD')} deg"],
            ["Eccentricity", str(orb.get("eccentricity", "0.0"))],
        ]
        if orb.get("period_min"):
            orbit_rows.append(["Orbital Period", f"{orb['period_min']} min"])
        if orb.get("orbits_per_day"):
            orbit_rows.append(["Orbits per Day", f"{orb['orbits_per_day']}"])
        if orb.get("eclipse_fraction"):
            orbit_rows.append(["Eclipse Fraction", f"{orb['eclipse_fraction']}"])
        if orb.get("eclipse_duration_min"):
            orbit_rows.append(["Eclipse Duration", f"{orb['eclipse_duration_min']} min"])
        if orb.get("velocity_ms"):
            orbit_rows.append(["Orbital Velocity", f"{orb['velocity_ms']} m/s"])
        if orb.get("design_lifetime_years"):
            orbit_rows.append(["Design Lifetime", f"{orb['design_lifetime_years']} years"])
        _add_table(doc, ["Parameter", "Value"], orbit_rows)
    else:
        doc.add_paragraph("[Orbit parameters to be defined]")

    # ── Section 3: Mission Phases ──
    _add_section(doc, "3", "Mission Phases")
    phases = ops.get("phases", [])
    if phases:
        _add_table(doc,
                   ["Phase", "Type", "Duration (days)", "Description"],
                   [
                       [
                           p.get("name", ""),
                           p.get("phase_type", ""),
                           str(p.get("duration_days", "")),
                           p.get("description", ""),
                       ]
                       for p in phases
                   ])
        # Add entry/exit criteria if available
        for p in phases:
            entry = p.get("entry_criteria", "")
            exit_c = p.get("exit_criteria", "")
            if entry or exit_c:
                _add_section(doc, f"3.{phases.index(p)+1}", p.get("name", "Phase"),
                             f"**Entry criteria:** {entry or 'TBD'}\n\n**Exit criteria:** {exit_c or 'TBD'}")
    else:
        _add_table(doc,
                   ["Phase", "Duration", "Description"],
                   [
                       ["LEOP", "3 days", "Launch and Early Orbit Phase: initial contact, detumble, solar array deployment"],
                       ["Commissioning", "30 days", "Platform and payload checkout, calibration, orbit verification"],
                       ["Nominal Operations", "Mission lifetime", "Science/service operations at full capability"],
                       ["Extended Operations", "TBD", "Reduced operations if consumables/orbit allow"],
                       ["Disposal", "TBD", "Passivation and deorbit per debris mitigation requirements"],
                   ])

    # ── Section 4: Operational Modes ──
    _add_section(doc, "4", "Operational Modes")
    modes = ops.get("modes", [])
    if modes:
        _add_table(doc,
                   ["Mode", "Type", "Power (W)", "Payload Active", "Data Rate (Mbps)",
                    "Pointing (deg)", "Duty Cycle (%)"],
                   [
                       [
                           m.get("name", ""),
                           m.get("mode_type", ""),
                           f"{m.get('power_w', 0):.1f}" if m.get("power_w") else "",
                           "Yes" if m.get("payload_active") else "No",
                           f"{m.get('data_rate_mbps', 0):.1f}" if m.get("data_rate_mbps") else "",
                           f"{m.get('pointing_requirement_deg', '')}" if m.get("pointing_requirement_deg") else "",
                           f"{m.get('duty_cycle_percent', '')}" if m.get("duty_cycle_percent") else "",
                       ]
                       for m in modes
                   ])
        # Mode descriptions
        for m in modes:
            desc = m.get("description", "")
            if desc:
                _render_bullet(doc, f"- **{m.get('name', '')}:** {desc}")
    else:
        _add_table(doc,
                   ["Mode", "Description", "Key Subsystems"],
                   [
                       ["Safe Mode", "Minimum power, sun-pointing, waiting for ground contact", "EPS, AOCS (coarse), TTC (beacon)"],
                       ["Nominal / Science", "Full payload operation, nadir-pointing", "All subsystems active"],
                       ["Downlink", "High-rate data downlink to ground station", "TTC (high power), OBC, AOCS"],
                       ["Eclipse", "Battery-powered, payload may be inactive", "EPS (battery), Thermal (heaters)"],
                   ])

    # ── Section 5: Ground Station Network ──
    _add_section(doc, "5", "Ground Station Network")
    gs_list = ops.get("ground_stations", [])
    # Also check elements for ground segment
    gs_elements = [e for e in elems if e.get("segment") == "ground"
                   and (e.get("performance") or {}).get("latitude")]
    if gs_list:
        _add_table(doc,
                   ["Station", "Type", "Latitude", "Longitude", "Antenna (m)",
                    "Bands", "Contact (min/day)"],
                   [
                       [
                           gs.get("name", ""),
                           gs.get("type", ""),
                           f"{gs.get('latitude_deg', ''):.1f}" if gs.get("latitude_deg") else "",
                           f"{gs.get('longitude_deg', ''):.1f}" if gs.get("longitude_deg") else "",
                           f"{gs.get('antenna_diameter_m', '')}" if gs.get("antenna_diameter_m") else "",
                           ", ".join(gs.get("frequency_bands", [])),
                           f"{gs.get('contact_time_per_day_min', '')}" if gs.get("contact_time_per_day_min") else "",
                       ]
                       for gs in gs_list
                   ])
    elif gs_elements:
        _add_table(doc,
                   ["Station", "Latitude", "Longitude", "Bands"],
                   [
                       [
                           e.get("name", ""),
                           str((e.get("performance") or {}).get("latitude", "")),
                           str((e.get("performance") or {}).get("longitude", "")),
                           ", ".join((e.get("performance") or {}).get("bands", [])),
                       ]
                       for e in gs_elements
                   ])
    else:
        doc.add_paragraph("[Ground station network to be defined]")

    # Operations concept
    ops_concept = ops.get("operations_concept", "")
    autonomy = ops.get("autonomy_level", "")
    if ops_concept or autonomy:
        _add_section(doc, "5.1", "Operations Concept")
        if ops_concept:
            _render_rich_content(doc, ops_concept)
        if autonomy:
            doc.add_paragraph(f"Autonomy level: {autonomy}")

    # ── Section 6: Data Flow ──
    _add_section(doc, "6", "Data Flow")
    pipeline = ops.get("data_pipeline", [])
    if pipeline:
        _add_table(doc,
                   ["Step", "Location", "Description", "Latency", "Data Level"],
                   [
                       [
                           dp.get("name", ""),
                           dp.get("location", ""),
                           dp.get("description", ""),
                           dp.get("latency", ""),
                           dp.get("data_level", ""),
                       ]
                       for dp in pipeline
                   ])
    else:
        _add_section(doc, "6.1", "Data Pipeline",
                     "- Instrument acquisition (onboard)\n"
                     "- Onboard storage (mass memory)\n"
                     "- Downlink via TTC subsystem to ground station\n"
                     "- Ground processing (L0 to L1/L2 products)\n"
                     "- Archive and distribution to end users")

    # Data rates from modes
    data_modes = [m for m in modes if m.get("data_rate_mbps", 0) > 0]
    if data_modes:
        _add_section(doc, "6.2", "Data Generation Rates")
        _add_table(doc,
                   ["Mode", "Data Rate (Mbps)"],
                   [
                       [m.get("name", ""), f"{m['data_rate_mbps']:.1f}"]
                       for m in data_modes
                   ])

    # ── Section 7: Anomaly Response ──
    _add_section(doc, "7", "Anomaly Response")
    _add_table(doc,
               ["Contingency", "Trigger", "Response", "Recovery"],
               [
                   ["Loss of attitude", "Attitude error exceeds threshold",
                    "Autonomous transition to Safe Mode, sun-pointing", "Ground diagnosis, mode recovery command"],
                   ["Power anomaly", "Battery SoC below safe threshold",
                    "Load shedding, payload shutdown", "Ground analysis, gradual load restoration"],
                   ["Communication loss", "Missed scheduled ground contact",
                    "Autonomous safe mode after timeout", "Next ground pass recovery attempt"],
                   ["Thermal exceedance", "Temperature sensor out of range",
                    "Heater activation or payload duty cycle reduction", "Ground thermal analysis, ops adjustment"],
                   ["Payload anomaly", "Payload health check failure",
                    "Payload power-off, continue platform operations", "Diagnostic data downlink, reconfiguration"],
               ])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_vp_docx(
    study_name: str = "",
    requirements: list[dict[str, Any]] | None = None,
) -> bytes:
    """Generate Verification Plan as .docx bytes."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)
    reqs = requirements or []

    _add_cover_page(doc, "Verification Plan", "ECSS-E-ST-10-02C Rev.1", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", f"Defines the verification approach for {study_name}.")
    _add_section(doc, "1.2", "Verification Methods",
                 "Analysis (A), Test (T), Review of Design (R), Inspection (I) per ECSS-E-ST-10-02C.")

    _add_section(doc, "2", "Verification Matrix")
    if reqs:
        _add_table(doc,
                   ["Req ID", "Requirement", "Method", "Phase", "Status"],
                   [[r.get("id", ""), r.get("text", "")[:60], r.get("verification_method", "A"),
                     "Phase B", "Planned"] for r in reqs[:50]])

    _add_section(doc, "3", "Environmental Test Programme")
    _add_section(doc, "3.1", "Vibration", "Sine + random per launch vehicle PUG. 3 axes, 1 min/axis.")
    _add_section(doc, "3.2", "Thermal Vacuum", "Qualification range +/-10C beyond operating. 4 cycles minimum.")
    _add_section(doc, "3.3", "EMC", "Per ECSS-E-ST-20-07C if required by launch provider.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_ts_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate Technical Specification (ECSS-E-ST-10C Annex B)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Technical Specification", "ECSS-E-ST-10C Annex B", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", "This document defines the technical requirements for the system and its elements.")
    _add_section(doc, "1.2", "Scope", f"Applicable to the {study_name} system at all levels of decomposition.")

    _add_section(doc, "2", "Applicable and Reference Documents")
    doc.add_paragraph("ECSS-E-ST-10C -- Space engineering -- System engineering general requirements")
    doc.add_paragraph("ECSS-Q-ST-40C -- Space product assurance -- Safety")

    _add_section(doc, "3", "System Description")
    mission_need = data.get("mission_need", {})
    doc.add_paragraph(mission_need.get("problem_statement", "To be defined"))

    _add_section(doc, "4", "Technical Requirements")
    reqs = data.get("requirements", [])
    if reqs:
        _add_table(doc, ["ID", "Type", "Requirement", "Threshold", "V-Method"],
                   [[r.get("id", ""), r.get("req_type", ""), r.get("text", "")[:80],
                     f"{r.get('operator', '')} {r.get('threshold', '')} {r.get('unit', '')}",
                     r.get("verification_method", "A")] for r in reqs[:80]])

    _add_section(doc, "5", "Interface Requirements")
    doc.add_paragraph("Interface requirements are defined in the IRD (separate document).")

    _add_section(doc, "6", "Design and Construction Standards")
    doc.add_paragraph("Per ECSS-Q-ST-70C for materials and processes. Workmanship per ECSS-Q-ST-70-08C.")

    _add_section(doc, "7", "Verification Requirements")
    doc.add_paragraph("Verification approach defined in the Verification Plan (VP). "
                      "All requirements shall be verified by at least one method: Analysis, Test, Inspection, or Demonstration.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_ird_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate Interface Requirements Document (NASA SEH Appendix L)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Interface Requirements Document", "NASA SEH Appendix L / ECSS-E-ST-10-24C", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", "This document defines the interfaces between mission segments and between subsystems.")
    _add_section(doc, "1.2", "Scope", "Covers mechanical, electrical, data, and RF interfaces.")

    _add_section(doc, "2", "Mission Segment Interfaces")
    _add_table(doc, ["Interface", "Segment A", "Segment B", "Type", "Description"],
               [["IF-001", "Space", "Ground", "RF", "S/X-band TM/TC link"],
                ["IF-002", "Space", "Ground", "RF", "Payload data downlink"],
                ["IF-003", "Ground", "User", "Data", "Data products via API/FTP"],
                ["IF-004", "Launch", "Space", "Mechanical", "Separation interface per PUG"],
                ["IF-005", "External", "Space", "RF", "GNSS signals for orbit determination"]])

    _add_section(doc, "3", "Subsystem Interfaces")
    _add_table(doc, ["Interface", "From", "To", "Type", "Protocol/Standard"],
               [["IF-101", "OBC", "EPS", "Electrical", "I2C power telemetry"],
                ["IF-102", "OBC", "TTC", "Data", "CCSDS TM/TC packets"],
                ["IF-103", "OBC", "Payload", "Data", "SpaceWire/LVDS"],
                ["IF-104", "EPS", "All", "Electrical", "28V regulated bus"],
                ["IF-105", "AOCS", "Payload", "Data", "Attitude quaternion @ 10Hz"]])

    _add_section(doc, "4", "Physical Interface Definitions")
    doc.add_paragraph("Connector definitions, pinouts, and harness routing to be defined during Phase C.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_semp_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate Systems Engineering Management Plan (NASA SEH Appendix J)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Systems Engineering Management Plan", "NASA SEH Appendix J / ECSS-M-ST-10C", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", "This SEMP defines the systems engineering approach, processes, and management controls.")
    _add_section(doc, "1.2", "Mission Overview")
    mission_need = data.get("mission_need", {})
    doc.add_paragraph(mission_need.get("problem_statement", "To be defined"))

    _add_section(doc, "2", "SE Organisation and Responsibilities")
    doc.add_paragraph("The SE team follows the CDF concurrent engineering model with 16 specialist positions.")
    _add_table(doc, ["Role", "Responsibility"],
               [["Systems Engineer", "Technical baseline, budgets, trade studies"],
                ["Mission Analyst", "Orbit, coverage, delta-V analysis"],
                ["Payload Lead", "Instrument performance, data products"],
                ["Power Engineer", "EPS sizing, energy balance"],
                ["AOCS Engineer", "Pointing, orbit maintenance"],
                ["Comms Engineer", "Link budget, TTC design"],
                ["Thermal Engineer", "Thermal control, analysis"],
                ["Structures Engineer", "Configuration, mass budget, AIT"]])

    _add_section(doc, "3", "Technical Processes")
    _add_section(doc, "3.1", "Requirements Engineering", "Suggest-then-approve pattern with SMART validation. Requirements linked to functions and traced to objectives.")
    _add_section(doc, "3.2", "Architecture Design", "Trade studies per subsystem with derived requirements. Options evaluated on mass, power, cost, TRL.")
    _add_section(doc, "3.3", "Interface Management", "Interfaces captured at segment and subsystem level. Managed via interface matrix.")
    _add_section(doc, "3.4", "Verification & Validation", "V&V matrix per ECSS-E-ST-10-02C. Methods: Analysis, Test, Inspection, Demonstration.")

    _add_section(doc, "4", "Technical Management Processes")
    _add_section(doc, "4.1", "Configuration Management", "Design baselines at SRR, PDR, CDR. Change control via board.")
    _add_section(doc, "4.2", "Risk Management", "5x5 risk matrix per ECSS-M-ST-80C. Continuous monitoring.")
    _add_section(doc, "4.3", "Technical Budgets", "Mass, power, link, pointing, delta-V, data, cost. Margins per ECSS design phase.")

    _add_section(doc, "5", "Reviews and Decision Gates")
    _add_table(doc, ["Review", "Phase Entry", "Key Deliverables"],
               [["MCR", "Pre-A \u2192 A", "Mission Need Statement, ConOps outline"],
                ["SRR", "A \u2192 B", "Requirements baseline, functional architecture"],
                ["PDR", "B \u2192 C", "Preliminary design, budgets with margin"],
                ["CDR", "C \u2192 D", "Detailed design, manufacturing drawings, test plan"],
                ["TRR", "D (test)", "Qualified hardware, test procedures"],
                ["FRR", "D (launch)", "Flight-ready spacecraft, ops procedures"]])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_rmp_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate Risk Management Plan (ECSS-M-ST-80C)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Risk Management Plan", "ECSS-M-ST-80C / NPR 8000.4", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", "This plan defines the risk management process for identification, assessment, mitigation, and monitoring of project risks.")

    _add_section(doc, "2", "Risk Management Process")
    _add_section(doc, "2.1", "Risk Identification", "Risks identified via brainstorming, checklists, heritage review, and technical analysis. "
                 "Each risk assigned unique ID and categorised (technical/programmatic/regulatory).")
    _add_section(doc, "2.2", "Risk Assessment", "Likelihood (1-5) x Consequence (1-5) = Risk Priority Number. "
                 "Green (<=4), Amber (5-9), Orange (10-15), Red (16-25).")
    _add_section(doc, "2.3", "Risk Mitigation", "For each risk: Accept / Mitigate / Transfer / Avoid. "
                 "Mitigation actions assigned to responsible position with deadline.")
    _add_section(doc, "2.4", "Risk Monitoring", "Risks reviewed at each design review (SRR, PDR, CDR). "
                 "Status tracked: Open / Mitigating / Closed / Accepted.")

    _add_section(doc, "3", "Risk Register")
    risks = data.get("risks", [])
    if risks:
        _add_table(doc, ["ID", "Risk", "L", "C", "Score", "Mitigation", "Status"],
                   [[r.get("id", ""), r.get("title", ""), str(r.get("likelihood", "")),
                     str(r.get("consequence", "")), str(r.get("likelihood", 0) * r.get("consequence", 0)),
                     r.get("mitigation", "")[:60], r.get("status", "")] for r in risks[:20]])
    else:
        doc.add_paragraph("Risk register to be populated during Phase A.")

    _add_section(doc, "4", "Risk Acceptance Criteria")
    doc.add_paragraph("Risks scoring <= 9 may be accepted by the Project Manager. "
                      "Risks scoring 10-15 require Systems Engineer review. "
                      "Risks scoring >= 16 require Programme Board approval.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_testplan_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate Test Plan (ECSS-E-ST-10-03C)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Test Plan", "ECSS-E-ST-10-03C / ECSS-E-ST-10-02C", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Introduction")
    _add_section(doc, "1.1", "Purpose", "This plan defines the test programme for qualification and acceptance of the spacecraft.")

    _add_section(doc, "2", "Test Philosophy")
    doc.add_paragraph("Protoflight approach: single model subjected to qualification levels with acceptance durations. "
                      "Per ECSS-E-ST-10-03C for CubeSat/small satellite class.")

    _add_section(doc, "3", "Test Programme")
    _add_table(doc, ["Test", "Level", "Standard", "Facility", "Duration"],
               [["Vibration (Sine)", "Qualification", "Launch vehicle PUG", "Shaker table", "3 axes, 2 min/axis"],
                ["Vibration (Random)", "Qualification", "Launch vehicle PUG", "Shaker table", "3 axes, 1 min/axis"],
                ["Thermal Vacuum", "Qualification", "+/-10 deg C beyond op range", "TVAC chamber", "4 cycles, 2h dwell"],
                ["Thermal Cycling", "Acceptance", "Operating range", "Chamber", "8 cycles, 1h dwell"],
                ["EMC/EMI", "Qualification", "ECSS-E-ST-20-07C", "Anechoic chamber", "1 day"],
                ["Functional", "Acceptance", "N/A", "Clean room", "Continuous"],
                ["Mass Properties", "Acceptance", "N/A", "Balance fixture", "1 day"],
                ["Deployment", "Qualification", "N/A", "0g simulator", "100+ actuations"]])

    _add_section(doc, "4", "Test Sequence")
    doc.add_paragraph("1. Incoming inspection -> 2. Functional test (baseline) -> 3. Vibration -> "
                      "4. Functional test -> 5. TVAC -> 6. Functional test -> 7. EMC -> "
                      "8. Final functional test -> 9. Mass properties -> 10. Delivery/integration.")

    _add_section(doc, "5", "Pass/Fail Criteria")
    doc.add_paragraph("All functional tests must pass with identical results pre- and post-environmental. "
                      "Any degradation triggers NCR (Non-Conformance Report) and MRB (Material Review Board).")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_regulatory_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate a regulatory filing template document."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Regulatory Filing Template", "SpaceCDF Auto-Generated", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Mission Overview")
    mission_need = data.get("mission_need", {})
    doc.add_paragraph(mission_need.get("problem_statement", "To be defined"))
    _add_section(doc, "2", "Orbital Parameters")
    doc.add_paragraph(f"This section should be completed with the specific orbital parameters for the {study_name} mission.")
    _add_section(doc, "3", "Frequency Coordination")
    doc.add_paragraph("Frequency bands, bandwidths, EIRP, and coordination requirements to be specified per ITU Radio Regulations.")
    _add_section(doc, "4", "Spacecraft Characteristics")
    doc.add_paragraph("Mass, dimensions, power, and technical characteristics as required by the regulatory authority.")
    _add_section(doc, "5", "Compliance Statement")
    doc.add_paragraph("Statement of compliance with applicable regulations, standards, and guidelines.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_review_package_docx(study_name: str, data: dict[str, Any]) -> bytes:
    """Generate a design review data package (SRR/PDR/CDR)."""
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    _add_cover_page(doc, "Design Review Data Package", "SpaceCDF Auto-Generated", study_name)
    _add_toc_placeholder(doc)
    _add_footer(doc, study_name)

    _add_section(doc, "1", "Review Objectives")
    doc.add_paragraph("This document compiles the design data package for review.")
    _add_section(doc, "2", "Mission Summary")
    mission_need = data.get("mission_need", {})
    doc.add_paragraph(mission_need.get("problem_statement", "To be defined"))
    _add_section(doc, "3", "Requirements Status")
    reqs = data.get("requirements", [])
    if reqs:
        _add_table(doc, ["ID", "Requirement", "Status"],
                   [[r.get("id", ""), r.get("text", "")[:60], r.get("status", "draft")] for r in reqs[:50]])
    _add_section(doc, "4", "Design Summary")
    doc.add_paragraph("Key design parameters and architecture decisions.")
    _add_section(doc, "5", "Budget Status")
    doc.add_paragraph("Mass, power, cost, link margin status vs allocation.")
    _add_section(doc, "6", "Risk Register")
    doc.add_paragraph("Current risk status and mitigations.")
    _add_section(doc, "7", "Open Actions")
    doc.add_paragraph("Actions from previous reviews and their status.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# Map of available document types
DOCX_GENERATORS = {
    "mrd": ("Mission Requirements Document", generate_mrd_docx),
    "conops": ("Concept of Operations", generate_conops_docx),
    "vp": ("Verification Plan", generate_vp_docx),
    "ts": ("Technical Specification", generate_ts_docx),
    "ird": ("Interface Requirements Document", generate_ird_docx),
    "semp": ("Systems Engineering Management Plan", generate_semp_docx),
    "rmp": ("Risk Management Plan", generate_rmp_docx),
    "testplan": ("Test Plan", generate_testplan_docx),
    # Regulatory filings
    "itu_api": ("ITU API Filing Template", _generate_regulatory_docx),
    "iaru": ("IARU Coordination Request", _generate_regulatory_docx),
    "rsssa": ("RSSSA Filing (Canada)", _generate_regulatory_docx),
    "export": ("Export Control Assessment", _generate_regulatory_docx),
    "copuos": ("UN Registration (COPUOS)", _generate_regulatory_docx),
    "eol": ("End-of-Life Compliance Report", _generate_regulatory_docx),
    # Review packages
    "srr": ("SRR Design Review Package", _generate_review_package_docx),
    "pdr": ("PDR Design Review Package", _generate_review_package_docx),
    "cdr": ("CDR Design Review Package", _generate_review_package_docx),
}
