"""SpaceCDF — Export API Router.

Endpoints for generating SMO simulator configs, design review
documents, and flight software architecture from a design study.
"""
from __future__ import annotations

import io
import math
import re

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import JSONResponse, StreamingResponse

from spacecdf_common.agents.base import DesignState
from spacecdf_common.models.study import MissionRequirements
from spacecdf_agents import DesignLoopOrchestrator

from .studies import get_study_store

router = APIRouter()


async def _run_design_for_study(study_id: str) -> tuple:
    """Run design loop for a study and return (state, requirements, result)."""
    studies = get_study_store()
    study = studies.get(study_id)
    if not study:
        raise HTTPException(status_code=404, detail=f"Study {study_id} not found")

    orchestrator = DesignLoopOrchestrator()
    orchestrator.initialise_agents()
    result = await orchestrator.run(study.requirements)

    if not result.final_state:
        raise HTTPException(status_code=500, detail="Design loop produced no final state")

    return result.final_state, study.requirements, result


# ─── Helper: derive CubeSat form factor from total mass ───

def _form_factor_from_mass(mass_kg: float) -> str:
    """Return CubeSat form factor string from total spacecraft mass."""
    if mass_kg < 2:
        return "1U"
    elif mass_kg < 4:
        return "2U"
    elif mass_kg < 6:
        return "3U"
    elif mass_kg < 12:
        return "6U"
    elif mass_kg < 24:
        return "12U"
    else:
        return "Custom"


# ─── Helper: read study orbit parameters ───

def _read_orbit_params(study) -> dict:
    """Extract orbit parameters from a study object, with safe defaults."""
    altitude_km = 500.0
    inclination_deg = 97.4
    orbit_type = "SSO"
    eccentricity = 0.0
    design_lifetime_years = 3.0
    try:
        orb = study.requirements.orbit
        altitude_km = orb.altitude_km or 500.0
        inclination_deg = orb.inclination_deg or 97.4
        orbit_type = orb.orbit_type.value if hasattr(orb.orbit_type, "value") else str(orb.orbit_type)
        eccentricity = orb.eccentricity or 0.0
    except Exception:
        pass
    try:
        design_lifetime_years = study.requirements.design_lifetime_years or 3.0
    except Exception:
        pass
    return {
        "altitude_km": altitude_km,
        "inclination_deg": inclination_deg,
        "orbit_type": orbit_type,
        "eccentricity": eccentricity,
        "design_lifetime_years": design_lifetime_years,
    }


# ─── Helpers: FSW enrichment from study data ───

def _derive_state_machine(study) -> dict:
    """Derive FSW state machine from ConOps modes."""
    states = []
    transitions = []
    conops = getattr(study, "conops", None)
    if not conops:
        return {"states": [], "transitions": []}

    modes = getattr(conops, "modes", [])
    for m in modes:
        mode_name = m.name if hasattr(m, "name") else str(m)
        mode_type = m.mode_type.value if hasattr(m, "mode_type") and hasattr(m.mode_type, "value") else ""
        states.append({
            "name": mode_name,
            "mode_type": mode_type,
            "is_critical": getattr(m, "is_critical", False),
            "power_w": getattr(m, "power_w", 0.0),
        })

    # Standard transitions: any mode can reach Safe, Safe can reach any
    mode_names = [s["name"] for s in states]
    safe_name = next((n for n in mode_names if "safe" in n.lower()), mode_names[0] if mode_names else "Safe")

    for name in mode_names:
        if name != safe_name:
            transitions.append({
                "from": name, "to": safe_name,
                "trigger": f"fault_detected OR cmd_go_safe",
                "priority": "critical",
            })
            transitions.append({
                "from": safe_name, "to": name,
                "trigger": f"cmd_enter_{re.sub(r'[^a-z0-9]', '_', name.lower())}",
                "priority": "nominal",
            })

    # Sequential transitions between non-safe modes
    non_safe = [n for n in mode_names if n != safe_name]
    for i in range(len(non_safe) - 1):
        transitions.append({
            "from": non_safe[i], "to": non_safe[i + 1],
            "trigger": f"mode_complete_{re.sub(r'[^a-z0-9]', '_', non_safe[i].lower())}",
            "priority": "nominal",
        })

    return {"states": states, "transitions": transitions}


def _derive_task_list(elements: list[dict]) -> list[dict]:
    """Derive FSW task list from element tree — one task per subsystem domain."""
    domain_priorities = {
        "eps": ("eps_manager", 2, 1000),
        "power": ("power_manager", 2, 1000),
        "aocs": ("aocs_controller", 1, 100),
        "ttc": ("ttc_handler", 3, 5000),
        "obc": ("obc_supervisor", 1, 500),
        "obdh": ("obdh_manager", 2, 1000),
        "payload": ("payload_controller", 5, 10000),
        "thermal": ("thermal_monitor", 6, 30000),
        "propulsion": ("propulsion_manager", 4, 5000),
        "structure": ("structure_monitor", 8, 60000),
        "structures": ("structure_monitor", 8, 60000),
    }
    seen_domains: set[str] = set()
    tasks = []
    for e in elements:
        domain = e.get("subsystem_domain", "")
        if not domain or domain in seen_domains:
            continue
        if e.get("segment") != "space":
            continue
        seen_domains.add(domain)
        task_name, priority, period = domain_priorities.get(
            domain, (f"{domain}_handler", 7, 10000)
        )
        tasks.append({
            "task_name": task_name,
            "subsystem_domain": domain,
            "priority": priority,
            "period_ms": period,
            "stack_size_bytes": 4096,
        })
    # Always include watchdog and housekeeping
    tasks.append({"task_name": "watchdog", "subsystem_domain": "system",
                  "priority": 0, "period_ms": 100, "stack_size_bytes": 1024})
    tasks.append({"task_name": "housekeeping_collector", "subsystem_domain": "system",
                  "priority": 9, "period_ms": 10000, "stack_size_bytes": 2048})
    return sorted(tasks, key=lambda t: t["priority"])


def _derive_memory_map(elements: list[dict]) -> dict:
    """Estimate OBC memory map from component count. 512KB typical CubeSat."""
    space_components = [e for e in elements if e.get("segment") == "space"
                        and e.get("element_type") == "component"]
    component_count = len(space_components)
    total_ram_kb = 512
    os_kb = 128
    fsw_kb = 256
    data_buffer_kb = 128
    # Scale FSW allocation slightly with component count
    if component_count > 10:
        fsw_kb = min(fsw_kb + (component_count - 10) * 4, 384)
        data_buffer_kb = total_ram_kb - os_kb - fsw_kb

    return {
        "total_ram_kb": total_ram_kb,
        "allocation": {
            "os_kernel": {"size_kb": os_kb, "description": "RTOS kernel + drivers"},
            "fsw_application": {"size_kb": fsw_kb, "description": "Flight software application code + state"},
            "data_buffer": {"size_kb": data_buffer_kb, "description": "Telemetry / science data ring buffer"},
        },
        "component_count": component_count,
        "estimated_task_count": component_count + 2,  # +watchdog +housekeeping
    }


def _derive_tm_tc(elements: list[dict], conops_modes: list) -> dict:
    """Derive TM/TC packet definitions from elements and ConOps modes."""
    # TM: one housekeeping packet per subsystem
    tm_packets = []
    seen_domains: set[str] = set()
    apid = 1
    for e in elements:
        domain = e.get("subsystem_domain", "")
        if not domain or domain in seen_domains or e.get("segment") != "space":
            continue
        seen_domains.add(domain)
        tm_packets.append({
            "apid": apid,
            "name": f"TM_{domain.upper()}_HK",
            "type": "housekeeping",
            "subsystem": domain,
            "rate_hz": 1.0 if domain in ("aocs", "eps", "obc") else 0.1,
            "size_bytes": 64,
        })
        apid += 1

    # TC: mode change commands + subsystem-specific
    tc_packets = []
    cmd_id = 1
    for mode in conops_modes:
        mode_name = mode.name if hasattr(mode, "name") else str(mode)
        tc_packets.append({
            "cmd_id": cmd_id,
            "name": f"TC_ENTER_{re.sub(r'[^A-Z0-9]', '_', mode_name.upper())}",
            "type": "mode_change",
            "target_mode": mode_name,
        })
        cmd_id += 1
    # Standard commands
    for cmd_name, cmd_type in [
        ("TC_RESET_OBC", "system"), ("TC_DEPLOY_ANTENNA", "mechanism"),
        ("TC_SET_PARAM", "configuration"), ("TC_REQUEST_TM", "data"),
    ]:
        tc_packets.append({"cmd_id": cmd_id, "name": cmd_name, "type": cmd_type})
        cmd_id += 1

    return {"tm_packets": tm_packets, "tc_packets": tc_packets}


def _derive_fault_handlers(elements: list[dict]) -> list[dict]:
    """Derive FDIR fault handlers from element domains."""
    handlers = []
    domains_present = {e.get("subsystem_domain", "") for e in elements
                       if e.get("segment") == "space"}

    fault_map = {
        "eps": {"fault": "power_low", "action": "safe_mode",
                "description": "Battery SoC below threshold — enter safe mode, shed non-essential loads"},
        "power": {"fault": "power_low", "action": "safe_mode",
                  "description": "Battery SoC below threshold — enter safe mode, shed non-essential loads"},
        "aocs": {"fault": "attitude_lost", "action": "sun_search",
                 "description": "Attitude determination lost — initiate sun search using coarse sensors"},
        "ttc": {"fault": "comms_lost", "action": "beacon_mode",
                "description": "No ground contact for N orbits — switch to low-rate beacon transmit"},
        "obc": {"fault": "obc_watchdog_timeout", "action": "warm_reset",
                "description": "OBC watchdog expired — perform warm reset and reload FSW from NVM"},
        "obdh": {"fault": "memory_fault", "action": "scrub_and_restart",
                 "description": "SEU detected in memory — scrub, verify checksum, restart task"},
        "thermal": {"fault": "over_temperature", "action": "load_shed",
                    "description": "Component temperature exceeds limit — reduce duty cycle, disable payload"},
        "payload": {"fault": "payload_anomaly", "action": "payload_safe",
                    "description": "Payload reports anomaly — disable payload, store diagnostic TM"},
        "propulsion": {"fault": "thruster_fault", "action": "inhibit_propulsion",
                       "description": "Thruster valve anomaly — close isolation valve, enter safe mode"},
    }

    for domain, handler in fault_map.items():
        if domain in domains_present:
            handlers.append({
                "subsystem": domain,
                "fault": handler["fault"],
                "action": handler["action"],
                "description": handler["description"],
                "severity": "critical" if domain in ("eps", "power", "aocs") else "major",
                "autonomous": True,
            })

    return handlers


# ─── Helpers: MBSE enrichment ───

def _derive_functional_architecture(study) -> list[dict]:
    """Derive functional architecture from study's functional decomposition or objectives."""
    func_decomp = getattr(study, "functional_decomposition", None)
    if func_decomp and func_decomp.functions:
        return [
            {
                "id": f.id,
                "name": f.name,
                "type": f.function_type.value if hasattr(f.function_type, "value") else str(f.function_type),
                "description": f.description,
                "parent_id": f.parent_function_id,
                "objective_ids": f.objective_ids,
                "derived_requirement_ids": f.derived_requirement_ids,
            }
            for f in func_decomp.functions
        ]
    # Fallback: generate generic functions from mission need objectives
    mission_need = getattr(study, "mission_need", None)
    objectives = getattr(mission_need, "objectives", []) if mission_need else []
    generic_functions = [
        {"id": "fn_acquire", "name": "Acquire data", "type": "observe",
         "description": "Collect mission data using payload instruments"},
        {"id": "fn_process", "name": "Process data", "type": "process",
         "description": "Onboard data processing and compression"},
        {"id": "fn_store", "name": "Store data", "type": "store",
         "description": "Buffer data between acquisition and downlink"},
        {"id": "fn_downlink", "name": "Downlink data", "type": "communicate",
         "description": "Transmit data to ground segment"},
        {"id": "fn_attitude", "name": "Maintain attitude", "type": "point",
         "description": "Attitude determination and control"},
        {"id": "fn_power", "name": "Generate power", "type": "power",
         "description": "Solar energy conversion, storage, and distribution"},
        {"id": "fn_thermal", "name": "Control thermal environment", "type": "protect",
         "description": "Maintain component temperatures within limits"},
        {"id": "fn_navigate", "name": "Navigate and manoeuvre", "type": "navigate",
         "description": "Orbit determination, station-keeping, collision avoidance"},
    ]
    # Link to objectives if present
    for i, obj in enumerate(objectives):
        obj_id = getattr(obj, "id", f"OBJ-{i+1}")
        if i < len(generic_functions):
            generic_functions[i]["objective_ids"] = [obj_id]
    return generic_functions


def _derive_logical_architecture(elements: list[dict], functions: list[dict]) -> list[dict]:
    """Map functions to systems/subsystems from element tree."""
    # Build domain → elements mapping
    domain_elements: dict[str, list[str]] = {}
    for e in elements:
        domain = e.get("subsystem_domain", "")
        if domain and e.get("segment") == "space":
            domain_elements.setdefault(domain, []).append(e.get("name", ""))

    function_to_domain = {
        "observe": ["payload"],
        "communicate": ["ttc"],
        "navigate": ["aocs", "propulsion"],
        "point": ["aocs"],
        "power": ["eps", "power"],
        "protect": ["thermal", "structure", "structures"],
        "process": ["obc", "obdh"],
        "store": ["obc", "obdh"],
        "propel": ["propulsion"],
        "support": ["structure", "structures"],
        "command": ["obc", "ttc"],
        "dispose": ["propulsion"],
    }

    mappings = []
    for func in functions:
        ftype = func.get("type", "")
        candidate_domains = function_to_domain.get(ftype, [])
        allocated_elements = []
        for d in candidate_domains:
            allocated_elements.extend(domain_elements.get(d, []))
        mappings.append({
            "function_id": func["id"],
            "function_name": func["name"],
            "allocated_domains": candidate_domains,
            "allocated_elements": allocated_elements,
        })
    return mappings


def _derive_physical_architecture(elements: list[dict]) -> list[dict]:
    """Return element tree with parent/child relationships."""
    element_map = {e.get("id"): e for e in elements if e.get("id")}
    tree = []
    for e in elements:
        children = [
            c.get("name", "") for c in elements
            if c.get("parent_element_id") == e.get("id")
        ]
        tree.append({
            "id": e.get("id", ""),
            "name": e.get("name", ""),
            "element_type": e.get("element_type", ""),
            "segment": e.get("segment", ""),
            "subsystem_domain": e.get("subsystem_domain", ""),
            "parent_id": e.get("parent_element_id", ""),
            "parent_name": element_map.get(e.get("parent_element_id", ""), {}).get("name", ""),
            "children": children,
            "mass_kg": e.get("mass_kg"),
            "power_avg_w": e.get("power_avg_w"),
            "quantity": e.get("quantity", 1),
        })
    return tree


def _derive_allocation_matrix(requirements: list[dict], elements: list[dict]) -> list[dict]:
    """Build requirements-to-elements allocation matrix from element_id FK."""
    element_map = {e.get("id"): e.get("name", "") for e in elements if e.get("id")}
    allocations = []
    for r in requirements:
        element_id = r.get("element_id", "")
        allocations.append({
            "requirement_code": r.get("code", r.get("id", "")),
            "requirement_text": r.get("text", ""),
            "level": r.get("level", ""),
            "element_id": element_id,
            "element_name": element_map.get(element_id, "unallocated"),
            "verification_method": r.get("verification_method", ""),
        })
    return allocations


def _derive_traceability(study, requirements: list[dict], elements: list[dict]) -> list[dict]:
    """Build objective -> requirement -> element -> verification chain."""
    mission_need = getattr(study, "mission_need", None)
    objectives = getattr(mission_need, "objectives", []) if mission_need else []

    element_map = {e.get("id"): e.get("name", "") for e in elements if e.get("id")}
    chains = []

    for obj in objectives:
        obj_id = getattr(obj, "id", "")
        obj_text = getattr(obj, "text", "")
        # Find requirements that trace to this objective (by parent or tag)
        linked_reqs = [
            r for r in requirements
            if r.get("parent_objective_id") == obj_id
            or obj_id in r.get("objective_ids", [])
            or obj_text.lower()[:20] in r.get("text", "").lower()
        ]
        if not linked_reqs:
            # Fallback: all system-level requirements
            linked_reqs = [r for r in requirements if r.get("level") in ("mission", "system")]

        for r in linked_reqs:
            element_id = r.get("element_id", "")
            chains.append({
                "objective_id": obj_id,
                "objective_text": obj_text,
                "requirement_code": r.get("code", r.get("id", "")),
                "requirement_text": r.get("text", ""),
                "element_id": element_id,
                "element_name": element_map.get(element_id, "unallocated"),
                "verification_method": r.get("verification_method", ""),
            })

    return chains


# ─── Helpers: SMO enrichment ───

def _derive_mission_profile(study) -> list[dict]:
    """Derive mission event timeline from ConOps phases."""
    conops = getattr(study, "conops", None)
    phases = getattr(conops, "phases", []) if conops else []

    # Standard event timeline
    events = [
        {"event": "separation", "time_offset": "T+0s",
         "description": "Spacecraft separation from launch vehicle"},
    ]

    cumulative_s = 0
    for p in phases:
        phase_name = p.name if hasattr(p, "name") else str(p)
        duration_days = getattr(p, "duration_days", 0) or 0
        phase_type = p.phase_type.value if hasattr(p, "phase_type") and hasattr(p.phase_type, "value") else ""

        # Map ConOps phases to simulation events
        if "leop" in phase_type.lower() or "leop" in phase_name.lower():
            events.append({"event": "detumble", "time_offset": "T+30min",
                           "description": "Detumble and initial stabilisation"})
            events.append({"event": "sun_acquisition", "time_offset": "T+1h",
                           "description": "Sun acquisition and solar array deployment"})
            events.append({"event": "antenna_deploy", "time_offset": "T+1.5h",
                           "description": "Antenna deployment and first beacon"})
            cumulative_s += duration_days * 86400
        elif "commissioning" in phase_type.lower() or "commissioning" in phase_name.lower():
            events.append({"event": "commissioning_start",
                           "time_offset": f"T+{cumulative_s/3600:.0f}h",
                           "description": "Begin commissioning — subsystem checkout"})
            cumulative_s += duration_days * 86400
        elif "nominal" in phase_type.lower() or "nominal" in phase_name.lower():
            events.append({"event": "nominal_ops",
                           "time_offset": f"T+{cumulative_s/3600:.0f}h",
                           "description": "Transition to nominal science operations"})
            cumulative_s += duration_days * 86400
        elif "disposal" in phase_type.lower() or "disposal" in phase_name.lower():
            events.append({"event": "end_of_life",
                           "time_offset": f"T+{cumulative_s/3600:.0f}h",
                           "description": "Begin disposal / passivation sequence"})

    # Estimate first ground pass from orbit
    events.append({"event": "first_ground_pass",
                   "time_offset": "T+~90min",
                   "description": "Estimated first ground station contact (orbit-dependent)"})

    return events


def _derive_perturbations(orbit_params: dict) -> dict:
    """Set perturbation model flags from orbit type."""
    orbit_type = orbit_params.get("orbit_type", "SSO").upper()
    altitude_km = orbit_params.get("altitude_km", 500.0)

    is_leo = orbit_type in ("LEO", "SSO") or altitude_km < 2000
    is_geo = orbit_type == "GEO" or (altitude_km > 35000 and altitude_km < 36000)
    is_deep = orbit_type in ("LUNAR", "INTERPLANETARY", "LAGRANGE")

    return {
        "J2": True,
        "atmospheric_drag": is_leo,
        "solar_radiation_pressure": True,
        "third_body_moon": is_geo or is_deep,
        "third_body_sun": is_geo or is_deep,
        "solid_earth_tides": False,
        "general_relativity": is_deep,
        "notes": (
            "LEO config: J2 + drag + SRP dominant"
            if is_leo else
            "GEO/deep-space config: third-body perturbations enabled"
            if is_geo or is_deep else
            "MEO config: J2 + SRP, drag negligible"
        ),
    }


def _derive_simulation_config(orbit_params: dict) -> dict:
    """Set simulation time step and duration from orbit."""
    altitude_km = orbit_params.get("altitude_km", 500.0)
    # Estimate orbital period
    mu = 3.986004418e5  # km^3/s^2
    r_earth = 6371.0
    a = r_earth + altitude_km
    period_s = 2 * math.pi * math.sqrt(a ** 3 / mu)

    return {
        "quick_sim": {
            "duration_s": round(period_s, 1),
            "duration_label": "1 orbit",
            "time_step_s": 10,
        },
        "full_sim": {
            "duration_s": 86400,
            "duration_label": "24 hours",
            "time_step_s": 10,
        },
        "long_sim": {
            "duration_s": 86400 * 7,
            "duration_label": "7 days",
            "time_step_s": 60,
        },
        "orbital_period_s": round(period_s, 1),
        "integrator": "RK45",
        "reference_frame": "J2000",
        "epoch": "2026-01-01T00:00:00Z",
    }


def _derive_output_config() -> dict:
    """Define simulation output variables."""
    return {
        "state_vector": {
            "position_eci_km": True,
            "velocity_eci_km_s": True,
            "position_ecef_km": True,
            "lat_lon_alt": True,
        },
        "attitude": {
            "quaternion": True,
            "euler_angles_deg": True,
            "angular_velocity_deg_s": True,
            "nadir_angle_deg": True,
        },
        "power": {
            "solar_flux_w_m2": True,
            "eclipse_flag": True,
            "battery_soc_percent": True,
            "power_generated_w": True,
            "power_consumed_w": True,
        },
        "thermal": {
            "solar_input_w": True,
            "albedo_input_w": True,
            "earth_ir_input_w": True,
            "spacecraft_temp_c": True,
        },
        "ground_contact": {
            "station_name": True,
            "elevation_deg": True,
            "in_contact": True,
        },
        "output_rate_hz": 0.1,
    }


def _derive_orbit_initial_conditions(orbit_params: dict) -> dict:
    """Compute initial Keplerian elements and Cartesian state from orbit params."""
    mu = 3.986004418e5  # km^3/s^2
    r_earth = 6371.0
    alt = orbit_params.get("altitude_km", 500.0)
    inc = orbit_params.get("inclination_deg", 97.4)
    ecc = orbit_params.get("eccentricity", 0.0)
    a = r_earth + alt
    period_s = 2 * math.pi * math.sqrt(a ** 3 / mu)
    v_circular = math.sqrt(mu / a)

    return {
        "keplerian": {
            "semi_major_axis_km": round(a, 3),
            "eccentricity": ecc,
            "inclination_deg": inc,
            "raan_deg": 0.0,
            "argument_of_perigee_deg": 0.0,
            "true_anomaly_deg": 0.0,
        },
        "cartesian_eci": {
            "x_km": round(a, 3),
            "y_km": 0.0,
            "z_km": 0.0,
            "vx_km_s": 0.0,
            "vy_km_s": round(v_circular * math.cos(math.radians(inc)), 4),
            "vz_km_s": round(v_circular * math.sin(math.radians(inc)), 4),
        },
        "derived": {
            "orbital_period_s": round(period_s, 1),
            "circular_velocity_km_s": round(v_circular, 4),
            "altitude_km": alt,
        },
    }


@router.post("/smo/{study_id}")
async def export_smo_config(study_id: str) -> JSONResponse:
    """Generate SMO simulator configuration files from a design study."""
    studies = get_study_store()
    study = studies.get(study_id)
    state, requirements, _ = await _run_design_for_study(study_id)

    from spacecdf_agents.exporters.smo.exporter import SMOExporter
    exporter = SMOExporter()
    export = exporter.export(state, requirements)

    # ── Enrichments: mission profile, perturbations, sim config ──
    orbit_params = _read_orbit_params(study) if study else {
        "altitude_km": 500.0, "inclination_deg": 97.4, "orbit_type": "SSO",
        "eccentricity": 0.0, "design_lifetime_years": 3.0,
    }

    mission_profile = _derive_mission_profile(study) if study else []
    perturbations = _derive_perturbations(orbit_params)
    simulation_config = _derive_simulation_config(orbit_params)
    output_config = _derive_output_config()
    orbit_initial_conditions = _derive_orbit_initial_conditions(orbit_params)

    return JSONResponse(content={
        "files": {k: v for k, v in export.files.items()},
        "is_valid": export.is_valid,
        "validation_errors": export.validation_errors,
        "validation_warnings": export.validation_warnings,
        "param_count": len(export.param_id_map),
        "file_count": len(export.files),
        "mission_profile": mission_profile,
        "perturbations": perturbations,
        "simulation_config": simulation_config,
        "output_config": output_config,
        "orbit_initial_conditions": orbit_initial_conditions,
    })


@router.get("/smo/{study_id}/validate")
async def validate_smo_export(study_id: str) -> dict:
    """Run SMO export validation only (no file generation)."""
    state, requirements, _ = await _run_design_for_study(study_id)

    from spacecdf_agents.exporters.smo.exporter import SMOExporter
    exporter = SMOExporter()
    export = exporter.export(state, requirements)

    return {
        "is_valid": export.is_valid,
        "errors": export.validation_errors,
        "warnings": export.validation_warnings,
    }


@router.post("/docs/{study_id}")
async def export_design_document(
    study_id: str,
    review: str = Query(default="srr", description="Review type: srr, pdr, cdr"),
    fmt: str = Query(default="zip", description="Output format: zip (default) or markdown"),
):
    """Generate a design review document from a design study.

    When ``fmt=zip`` (default) streams a zip archive containing Markdown,
    Word, and Excel artefacts. When ``fmt=markdown``, returns the Markdown
    body as JSON for backwards compatibility.
    """
    state, requirements, result = await _run_design_for_study(study_id)

    from spacecdf_agents.exporters.docs.generator import DocumentGenerator
    generator = DocumentGenerator()
    review_lower = (review or "srr").lower()
    if review_lower not in ("srr", "pdr", "cdr"):
        review_lower = "srr"

    if fmt.lower() == "markdown":
        doc = generator.generate(state, requirements, result, review_type=review_lower)
        return JSONResponse(content={
            "review_type": review_lower,
            "document_markdown": doc,
            "sections": doc.count("\n## "),
        })

    zip_bytes = generator.generate_bundle(
        state, requirements, result,
        review_type=review_lower,
        study_name=study_id,
    )
    filename = f"{review_lower}_export.zip"
    return StreamingResponse(
        io.BytesIO(zip_bytes),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/fsw/{study_id}")
async def export_fsw_architecture(study_id: str) -> JSONResponse:
    """Generate flight software architecture scaffolding."""
    studies = get_study_store()
    study = studies.get(study_id)
    state, requirements, _ = await _run_design_for_study(study_id)

    # ── Gather element and ConOps data for enrichment ──
    elements: list[dict] = []
    conops_modes: list = []
    try:
        from .elements import _elements
        elements = [e for e in _elements.values()
                    if e.get("study_id") == study_id and not e.get("deleted_at")]
    except Exception:
        pass
    if study:
        conops = getattr(study, "conops", None)
        conops_modes = getattr(conops, "modes", []) if conops else []

    # ── Derive FSW enrichments ──
    fsw_state_machine = _derive_state_machine(study) if study else {"states": [], "transitions": []}
    fsw_task_list = _derive_task_list(elements)
    fsw_memory_map = _derive_memory_map(elements)
    fsw_tm_tc = _derive_tm_tc(elements, conops_modes)
    fsw_fault_handlers = _derive_fault_handlers(elements)

    from spacecdf_agents.exporters.fsw.generator import FSWGenerator
    generator = FSWGenerator()
    files = generator.generate(state, requirements)

    return JSONResponse(content={
        "files": files,
        "file_count": len(files),
        "state_machine": fsw_state_machine,
        "task_list": fsw_task_list,
        "memory_map": fsw_memory_map,
        "tm_tc_definitions": fsw_tm_tc,
        "fault_handlers": fsw_fault_handlers,
    })


@router.post("/mbse/{study_id}")
async def export_mbse(study_id: str) -> JSONResponse:
    """Generate an ECSS-E-TM-10-25A-style MBSE JSON export.

    Research-credibility artefact: SysML-like model with blocks, parameters,
    requirements, traceability links, and applicable ECSS standards.
    Suitable for import into Cameo / Capella via a downstream converter and
    for archival in version control (diff-friendly JSON).
    """
    studies = get_study_store()
    study = studies.get(study_id)
    if not study:
        raise HTTPException(status_code=404, detail=f"Study {study_id} not found")

    state, requirements, _ = await _run_design_for_study(study_id)

    # Auto-generate requirements from the mission baseline when the study
    # hasn't been augmented with explicit ones.
    from spacecdf_common.models.requirements import generate_requirements
    req_objects = generate_requirements(requirements.model_dump())

    # Pull applicable ECSS from the originating template when present
    applicable_ecss: list[str] = []
    notes = getattr(study, "notes", "") or ""
    if "Seeded from template:" in notes:
        tmpl_id = notes.split("Seeded from template:")[1].strip().split()[0]
        try:
            from ..services.template_library import get_template
            tmpl = get_template(tmpl_id)
            if tmpl:
                applicable_ecss = list(tmpl.applicable_ecss)
        except Exception:
            pass

    from spacecdf_agents.exporters.mbse import generate_mbse_export

    phase = study.phase.value if hasattr(study.phase, "value") else str(study.phase)
    export = generate_mbse_export(
        study_id=study.id,
        study_name=study.name,
        phase=phase,
        parameters=state.parameters,
        requirements=req_objects,
        applicable_standards=applicable_ecss,
        notes=notes,
    )

    # ── MBSE Enrichments: functional / logical / physical arch, allocations, traceability ──
    elements: list[dict] = []
    interfaces: list[dict] = []
    req_dicts: list[dict] = []
    try:
        from .elements import _elements, _interfaces
        elements = [e for e in _elements.values()
                    if e.get("study_id") == study_id and not e.get("deleted_at")]
        interfaces = [
            i for i in _interfaces.values()
            if not i.get("deleted_at") and (
                i.get("study_id") == study_id
                or i.get("source_element_id") in {e.get("id") for e in elements}
                or i.get("target_element_id") in {e.get("id") for e in elements}
            )
        ]
    except Exception:
        pass
    try:
        from .requirements import _requirements
        req_dicts = [r for r in _requirements.values()
                     if r.get("study_id") == study_id and r.get("status") != "retired"]
    except Exception:
        pass

    functional_architecture = _derive_functional_architecture(study)
    logical_architecture = _derive_logical_architecture(elements, functional_architecture)
    physical_architecture = _derive_physical_architecture(elements)

    # Interface definitions from the interface store
    interface_definitions = []
    element_map = {e.get("id"): e.get("name", "") for e in elements if e.get("id")}
    for iface in interfaces:
        interface_definitions.append({
            "id": iface.get("id", ""),
            "source_element_id": iface.get("source_element_id", ""),
            "source_name": element_map.get(iface.get("source_element_id", ""),
                                           iface.get("source_name", "")),
            "target_element_id": iface.get("target_element_id", ""),
            "target_name": element_map.get(iface.get("target_element_id", ""),
                                           iface.get("target_name", "")),
            "interface_type": iface.get("interface_type", iface.get("type", "")),
            "properties": iface.get("properties", {}),
            "direction": iface.get("direction", "bidirectional"),
        })

    # Also include subsystem-level interfaces from study's InterfaceMatrix
    interface_matrix = getattr(study, "interface_matrix", None)
    if interface_matrix:
        for si in getattr(interface_matrix, "subsystem_interfaces", []):
            interface_definitions.append({
                "id": f"{si.subsystem_a}_{si.subsystem_b}",
                "source_name": si.subsystem_a,
                "target_name": si.subsystem_b,
                "interface_type": [it.value if hasattr(it, "value") else str(it)
                                   for it in getattr(si, "interface_types", [])],
                "status": si.status.value if hasattr(si.status, "value") else str(si.status),
                "criticality": getattr(si, "criticality", "standard"),
                "description": getattr(si, "description", ""),
            })

    allocation_matrix = _derive_allocation_matrix(req_dicts, elements)
    traceability = _derive_traceability(study, req_dicts, elements)

    export["functional_architecture"] = functional_architecture
    export["logical_architecture"] = logical_architecture
    export["physical_architecture"] = physical_architecture
    export["interface_definitions"] = interface_definitions
    export["allocation_matrix"] = allocation_matrix
    export["traceability"] = traceability

    return JSONResponse(content=export)


# --- Word Document Generation ---

@router.post("/docx/{doc_type}")
async def generate_docx(doc_type: str, study_id: str | None = None):
    """Generate an editable Word (.docx) document.

    Available types: mrd, conops, vp
    Returns a downloadable .docx file.
    """
    from ..services.docx_generator import DOCX_GENERATORS
    from .studies import get_study_store

    if doc_type not in DOCX_GENERATORS:
        raise HTTPException(400, f"Unknown doc type: {doc_type}. Available: {list(DOCX_GENERATORS.keys())}")

    name, gen_fn = DOCX_GENERATORS[doc_type]

    # Get study data if available
    study_name = "Unnamed Mission"
    mission_need: dict = {}
    requirements: list = []
    conops: dict = {}
    orbit_params: dict = {}
    elements: list = []
    interfaces: list = []

    if study_id:
        store = get_study_store()
        study = store.get(study_id)
        if study:
            study_name = study.name

            # ── Full mission_need dict ──
            if hasattr(study, 'mission_need') and study.mission_need:
                mn = study.mission_need
                mission_need = {
                    "problem_statement": getattr(mn, "problem_statement", ""),
                    "operational_context": getattr(mn, "operational_context", ""),
                    "selection_rationale": getattr(mn, "selection_rationale", ""),
                    "conops_summary": getattr(mn, "conops_summary", ""),
                    "objectives": [
                        {
                            "id": getattr(o, "id", ""),
                            "text": o.text,
                            "priority": o.priority.value if hasattr(o.priority, "value") else str(o.priority),
                            "type": o.type.value if hasattr(o.type, "value") else str(getattr(o, "type", "")),
                            "measurable_criterion": getattr(o, "measurable_criterion", ""),
                            "status": getattr(o, "status", "proposed"),
                        }
                        for o in getattr(mn, "objectives", [])
                    ],
                    "stakeholders": [
                        {
                            "name": s.name,
                            "role": s.role.value if hasattr(s.role, "value") else str(getattr(s, "role", "")),
                            "needs": getattr(s, "needs", []),
                            "priority": getattr(s, "priority", "primary"),
                        }
                        for s in getattr(mn, "stakeholders", [])
                    ],
                    "alternatives": [
                        {
                            "id": getattr(a, "id", ""),
                            "name": a.name,
                            "type": a.type.value if hasattr(a.type, "value") else str(getattr(a, "type", "")),
                            "description": getattr(a, "description", ""),
                            "pros": getattr(a, "pros", []),
                            "cons": getattr(a, "cons", []),
                            "feasibility_score": getattr(a, "feasibility_score", 0.0),
                            "decision": a.decision.value if hasattr(a.decision, "value") else str(getattr(a, "decision", "")),
                        }
                        for a in getattr(mn, "alternatives_considered", [])
                    ],
                    "selected_alternative_id": getattr(mn, "selected_alternative_id", None),
                }

            # ── Requirements from the requirements store ──
            try:
                from .requirements import _requirements
                requirements = [
                    r for r in _requirements.values()
                    if r.get("study_id") == study_id and r.get("status") != "retired"
                ]
            except Exception:
                pass

            # ── Orbit parameters ──
            orbit_params = _read_orbit_params(study)

            # ── Computed orbit values ──
            try:
                from spacecdf_common.physics.orbit import compute_orbit_params
                op = compute_orbit_params(
                    altitude_km=orbit_params["altitude_km"],
                    inclination_deg=orbit_params["inclination_deg"],
                    eccentricity=orbit_params.get("eccentricity", 0.0),
                )
                orbit_params.update({
                    "period_min": round(op.period_min, 2),
                    "velocity_ms": round(op.velocity_ms, 1),
                    "eclipse_fraction": round(op.eclipse_fraction, 3),
                    "eclipse_duration_min": round(getattr(op, "eclipse_duration_min", 0), 2),
                    "orbits_per_day": round(op.orbits_per_day, 1),
                })
            except Exception:
                pass

            # ── ConOps from study ──
            if hasattr(study, 'conops') and study.conops:
                cop = study.conops
                conops = {
                    "phases": [
                        {
                            "name": p.name, "phase_type": p.phase_type.value if hasattr(p.phase_type, "value") else str(p.phase_type),
                            "duration_days": p.duration_days, "description": p.description,
                            "entry_criteria": getattr(p, "entry_criteria", ""),
                            "exit_criteria": getattr(p, "exit_criteria", ""),
                        }
                        for p in getattr(cop, "phases", [])
                    ],
                    "modes": [
                        {
                            "name": m.name, "mode_type": m.mode_type.value if hasattr(m.mode_type, "value") else str(m.mode_type),
                            "description": getattr(m, "description", ""),
                            "power_w": m.power_w, "payload_active": m.payload_active,
                            "data_rate_mbps": m.data_rate_mbps,
                            "pointing_requirement_deg": m.pointing_requirement_deg,
                            "nadir_pointing": m.nadir_pointing,
                            "duty_cycle_percent": m.duty_cycle_percent,
                            "is_critical": m.is_critical,
                        }
                        for m in getattr(cop, "modes", [])
                    ],
                    "ground_stations": [
                        {
                            "name": gs.name,
                            "type": gs.type.value if hasattr(gs.type, "value") else str(gs.type),
                            "latitude_deg": gs.latitude_deg, "longitude_deg": gs.longitude_deg,
                            "antenna_diameter_m": gs.antenna_diameter_m,
                            "frequency_bands": gs.frequency_bands,
                            "contact_time_per_day_min": gs.contact_time_per_day_min,
                        }
                        for gs in getattr(cop, "ground_stations", [])
                    ],
                    "data_pipeline": [
                        {
                            "name": dp.name, "location": dp.location,
                            "description": dp.description, "latency": dp.latency,
                            "data_level": dp.data_level,
                        }
                        for dp in getattr(cop, "data_pipeline", [])
                    ],
                    "autonomy_level": getattr(cop, "autonomy_level", ""),
                    "operations_concept": getattr(cop, "operations_concept", ""),
                    "summary": getattr(cop, "summary", ""),
                }

            # ── Elements from the element store ──
            try:
                from .elements import _elements, _interfaces
                elements = [
                    e for e in _elements.values()
                    if e.get("study_id") == study_id and not e.get("deleted_at")
                ]
                interfaces = [
                    i for i in _interfaces.values()
                    if not i.get("deleted_at") and (
                        i.get("study_id") == study_id
                        or i.get("source_element_id") in {e.get("id") for e in elements}
                        or i.get("target_element_id") in {e.get("id") for e in elements}
                    )
                ]
            except Exception:
                pass

    # Build a universal data dict for generators that take (study_name, data)
    data = {
        "mission_need": mission_need,
        "requirements": requirements,
        "orbit": orbit_params,
        "conops": conops,
        "elements": elements,
        "interfaces": interfaces,
    }

    # Generate the document
    kwargs: dict = {"study_name": study_name}
    if doc_type in ("mrd",):
        kwargs["mission_need"] = mission_need
        kwargs["requirements"] = requirements
        kwargs["orbit"] = orbit_params
        kwargs["elements"] = elements
    elif doc_type in ("conops",):
        kwargs["mission_need"] = mission_need
        kwargs["conops"] = conops
        kwargs["orbit"] = orbit_params
        kwargs["elements"] = elements
    elif doc_type in ("vp",):
        kwargs["requirements"] = requirements
    elif doc_type in ("ts", "ird", "semp", "rmp", "testplan",
                      "itu_api", "iaru", "rsssa", "export", "copuos", "eol",
                      "srr", "pdr", "cdr"):
        kwargs["data"] = data

    docx_bytes = gen_fn(**kwargs)

    filename = f"SpaceCDF_{name.replace(' ', '_')}_{study_name.replace(' ', '_')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── New branded export endpoints ───

@router.post("/launch-icd/{study_id}")
async def export_launch_icd(
    study_id: str,
    fmt: str = Query(default="json", description="Output format: json, docx"),
) -> JSONResponse:
    """Generate Launch Interface Control Document data."""
    from ..services.branding import get_branding
    from .elements import _elements

    # Read study for orbit params and requirements
    studies = get_study_store()
    study = studies.get(study_id)
    orbit_params = _read_orbit_params(study) if study else {
        "altitude_km": 500.0, "inclination_deg": 97.4, "orbit_type": "SSO",
        "eccentricity": 0.0, "design_lifetime_years": 3.0,
    }

    elements = [e for e in _elements.values() if e.get("study_id") == study_id and not e.get("deleted_at")]
    spacecraft = [e for e in elements if e.get("segment") == "space" and e.get("element_type") in ("system", "segment")]
    components = [e for e in elements if e.get("element_type") == "component"]

    total_mass = sum((e.get("mass_kg") or 0) * (e.get("quantity", 1)) for e in elements if e.get("segment") == "space")
    comp_mass = round(sum((c.get("mass_kg") or 0) * c.get("quantity", 1) for c in components), 2)

    # Derive form factor from total mass
    form_factor = _form_factor_from_mass(total_mass)

    # Build component BOM table
    component_bom = []
    for c in components:
        component_bom.append({
            "name": c["name"],
            "subsystem": c.get("subsystem_domain", ""),
            "mass_kg": c.get("mass_kg"),
            "power_avg_w": c.get("power_avg_w"),
            "quantity": c.get("quantity", 1),
        })

    # Compute orbit parameters using physics module
    orbit_computed: dict = {}
    try:
        from spacecdf_common.physics.orbit import compute_orbit_params
        op = compute_orbit_params(
            altitude_km=orbit_params["altitude_km"],
            inclination_deg=orbit_params["inclination_deg"],
            eccentricity=orbit_params["eccentricity"],
        )
        orbit_computed = {
            "period_min": round(op.period_min, 2),
            "velocity_ms": round(op.velocity_ms, 1),
            "eclipse_fraction": round(op.eclipse_fraction, 3),
            "orbits_per_day": round(op.orbits_per_day, 1),
            "footprint_radius_km": round(op.footprint_radius_km, 1),
        }
    except Exception:
        pass

    b = get_branding()
    data = {
        "document": "Launch Interface Control Document",
        "branding": {"university": b.university, "department": b.department, "classification": b.classification},
        "study_id": study_id,
        "orbit": {
            "altitude_km": orbit_params["altitude_km"],
            "inclination_deg": orbit_params["inclination_deg"],
            "orbit_type": orbit_params["orbit_type"],
            "eccentricity": orbit_params["eccentricity"],
            **orbit_computed,
        },
        "spacecraft": {
            "total_mass_kg": round(total_mass, 2),
            "form_factor": form_factor,
            "systems": [{"name": s["name"], "mass_kg": s.get("mass_kg"), "quantity": s.get("quantity", 1)} for s in spacecraft],
        },
        "mechanical_interface": {
            "deployer_type": "Standard CubeSat deployer (ISIPOD / EXApod)",
            "rail_spec": "PC/104 compliant rails",
            "protrusion_limit_mm": 6.5,
            "cg_offset_limit_mm": 20,
        },
        "electrical_interface": {
            "inhibit_switches": 3,
            "battery_state": "Charged, RBF pin installed",
            "max_voltage_v": 8.4,
            "umbilical": "None (autonomous activation)",
        },
        "environmental": {
            "vibration": "Qualification: 14.1 grms (20-2000 Hz random)",
            "shock": "1500g SRS at 1000 Hz",
            "thermal_range_c": [-40, 60],
            "depressurization_rate": "< 5 kPa/s",
        },
        "component_bom": component_bom,
        "components_summary": {
            "total_components": len(components),
            "total_mass_kg": comp_mass,
        },
    }

    if fmt.lower() == "docx":
        from ..services.branding import create_branded_docx
        doc = create_branded_docx("Launch Interface Control Document", f"Study {study_id}")
        if doc is None:
            raise HTTPException(status_code=500, detail="python-docx not available")

        doc.add_heading("1. Orbit Parameters", level=1)
        doc.add_paragraph(f"Altitude: {orbit_params['altitude_km']} km")
        doc.add_paragraph(f"Inclination: {orbit_params['inclination_deg']} deg")
        doc.add_paragraph(f"Orbit type: {orbit_params['orbit_type']}")
        doc.add_paragraph(f"Eccentricity: {orbit_params['eccentricity']}")
        if orbit_computed:
            doc.add_paragraph(f"Period: {orbit_computed.get('period_min', 'TBD')} min")
            doc.add_paragraph(f"Velocity: {orbit_computed.get('velocity_ms', 'TBD')} m/s")
            doc.add_paragraph(f"Orbits per day: {orbit_computed.get('orbits_per_day', 'TBD')}")

        doc.add_heading("2. Spacecraft Description", level=1)
        doc.add_paragraph(f"Total mass: {round(total_mass, 2)} kg")
        doc.add_paragraph(f"Form factor: {form_factor}")
        if spacecraft:
            t = doc.add_table(rows=1, cols=3, style="Light List Accent 1")
            t.rows[0].cells[0].text = "System"
            t.rows[0].cells[1].text = "Mass (kg)"
            t.rows[0].cells[2].text = "Qty"
            for s in spacecraft:
                row = t.add_row()
                row.cells[0].text = s["name"]
                row.cells[1].text = str(s.get("mass_kg") or "TBD")
                row.cells[2].text = str(s.get("quantity", 1))

        doc.add_heading("3. Component Bill of Materials", level=1)
        if component_bom:
            t = doc.add_table(rows=1, cols=5, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Component"
            t.rows[0].cells[1].text = "Subsystem"
            t.rows[0].cells[2].text = "Mass (kg)"
            t.rows[0].cells[3].text = "Power (W)"
            t.rows[0].cells[4].text = "Qty"
            for c in component_bom:
                row = t.add_row()
                row.cells[0].text = c["name"]
                row.cells[1].text = c.get("subsystem", "")
                row.cells[2].text = str(c.get("mass_kg") or "TBD")
                row.cells[3].text = str(c.get("power_avg_w") or "TBD")
                row.cells[4].text = str(c.get("quantity", 1))
        else:
            doc.add_paragraph("No components defined.")

        doc.add_heading("4. Mechanical Interface", level=1)
        mi = data["mechanical_interface"]
        doc.add_paragraph(f"Deployer type: {mi['deployer_type']}")
        doc.add_paragraph(f"Rail spec: {mi['rail_spec']}")
        doc.add_paragraph(f"Protrusion limit: {mi['protrusion_limit_mm']} mm")
        doc.add_paragraph(f"CG offset limit: {mi['cg_offset_limit_mm']} mm")

        doc.add_heading("5. Electrical Interface", level=1)
        ei = data["electrical_interface"]
        doc.add_paragraph(f"Inhibit switches: {ei['inhibit_switches']}")
        doc.add_paragraph(f"Battery state: {ei['battery_state']}")
        doc.add_paragraph(f"Max voltage: {ei['max_voltage_v']} V")
        doc.add_paragraph(f"Umbilical: {ei['umbilical']}")

        doc.add_heading("6. Environmental Requirements", level=1)
        env = data["environmental"]
        doc.add_paragraph(f"Vibration: {env['vibration']}")
        doc.add_paragraph(f"Shock: {env['shock']}")
        doc.add_paragraph(f"Thermal range: {env['thermal_range_c'][0]} to {env['thermal_range_c'][1]} C")
        doc.add_paragraph(f"Depressurization rate: {env['depressurization_rate']}")

        doc.add_heading("7. Components Summary", level=1)
        doc.add_paragraph(f"Total components: {len(components)}")
        doc.add_paragraph(f"Total component mass: {comp_mass} kg")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Launch_ICD_{study_id}.docx"'},
        )

    return JSONResponse(content=data)


@router.post("/rsssa/{study_id}")
async def export_rsssa(
    study_id: str,
    fmt: str = Query(default="json", description="Output format: json, docx"),
) -> JSONResponse:
    """Generate RSSSA (Remote Sensing Space Systems Act) filing data."""
    from ..services.branding import get_branding
    from .elements import _elements

    # Read study for orbit params and requirements
    studies = get_study_store()
    study = studies.get(study_id)
    orbit_params = _read_orbit_params(study) if study else {
        "altitude_km": 500.0, "inclination_deg": 97.4, "orbit_type": "SSO",
        "eccentricity": 0.0, "design_lifetime_years": 3.0,
    }

    elements = [e for e in _elements.values() if e.get("study_id") == study_id and not e.get("deleted_at")]
    gs_elements = [e for e in elements if e.get("segment") == "ground" and (e.get("performance") or {}).get("latitude")]
    ttc_elements = [e for e in elements if e.get("subsystem_domain") == "ttc"]
    spacecraft = [e for e in elements if e.get("segment") == "space" and e.get("element_type") == "system"]
    payload_elements = [e for e in elements if e.get("subsystem_domain") == "payload"]

    total_mass = sum((e.get("mass_kg") or 0) * e.get("quantity", 1) for e in elements if e.get("segment") == "space")
    form_factor = _form_factor_from_mass(total_mass)

    # Build ground station list for regulatory template
    gs_list = [
        {"name": e["name"], "latitude": e["performance"]["latitude"],
         "longitude": e["performance"]["longitude"],
         "bands": e["performance"].get("bands", [])}
        for e in gs_elements
    ]

    # Extract payload pointing accuracy for GSD estimate
    pointing_accuracy_deg = None
    try:
        pointing_accuracy_deg = study.requirements.payloads[0].pointing_accuracy_deg if study else None
    except Exception:
        pass
    # Also check payload element performance
    if pointing_accuracy_deg is None:
        for p in payload_elements:
            perf = p.get("performance") or {}
            if perf.get("pointing_accuracy_deg"):
                pointing_accuracy_deg = perf["pointing_accuracy_deg"]
                break

    # Call generate_rsssa_template with actual study data
    filing_data: dict = {}
    try:
        from ..services.regulatory import generate_rsssa_template
        filing_data = generate_rsssa_template(
            study_name=study.name if study else "",
            operator_name=get_branding().university,
            orbit_altitude_km=orbit_params["altitude_km"],
            orbit_inclination_deg=orbit_params["inclination_deg"],
            orbit_type=orbit_params["orbit_type"],
            orbit_eccentricity=orbit_params["eccentricity"],
            mass_kg=total_mass or None,
            dimensions=form_factor,
            design_lifetime_years=orbit_params["design_lifetime_years"],
            pointing_accuracy_deg=pointing_accuracy_deg,
            ground_stations=gs_list if gs_list else None,
        )
    except Exception:
        pass

    # Run orbital lifetime estimate
    orbital_lifetime_years: float | None = None
    deorbit_compliant_25yr: bool | None = None
    try:
        from spacecdf_common.physics.debris import estimate_orbital_lifetime, estimate_cubesat_cross_section
        cross_section = estimate_cubesat_cross_section(form_factor) if form_factor != "Custom" else 0.03
        a_over_m = cross_section / max(total_mass, 0.1)
        orbital_lifetime_years = round(estimate_orbital_lifetime(
            altitude_km=orbit_params["altitude_km"],
            area_to_mass_ratio_m2_kg=a_over_m,
        ), 2)
        deorbit_compliant_25yr = orbital_lifetime_years <= 25.0
    except Exception:
        pass

    # Compute GSD estimate if pointing accuracy available
    gsd_estimate_m: float | None = None
    if pointing_accuracy_deg is not None:
        try:
            import math
            # Diffraction-limited GSD ≈ altitude * tan(pointing_accuracy)
            # Simplified: GSD ~ altitude_m * pointing_accuracy_rad
            alt_m = orbit_params["altitude_km"] * 1000
            gsd_estimate_m = round(alt_m * math.radians(pointing_accuracy_deg), 2)
        except Exception:
            pass

    # Compute orbit parameters
    orbit_computed: dict = {}
    try:
        from spacecdf_common.physics.orbit import compute_orbit_params
        op = compute_orbit_params(
            altitude_km=orbit_params["altitude_km"],
            inclination_deg=orbit_params["inclination_deg"],
            eccentricity=orbit_params["eccentricity"],
        )
        orbit_computed = {
            "period_min": round(op.period_min, 2),
            "eclipse_fraction": round(op.eclipse_fraction, 3),
            "orbits_per_day": round(op.orbits_per_day, 1),
        }
    except Exception:
        pass

    b = get_branding()
    data = {
        "document": "RSSSA Licence Application Data",
        "branding": {"university": b.university, "department": b.department, "classification": b.classification},
        "study_id": study_id,
        "applicant": {
            "name": b.university,
            "department": b.department,
            "country": "Canada",
        },
        "orbit": {
            "altitude_km": orbit_params["altitude_km"],
            "inclination_deg": orbit_params["inclination_deg"],
            "orbit_type": orbit_params["orbit_type"],
            "eccentricity": orbit_params["eccentricity"],
            "design_lifetime_years": orbit_params["design_lifetime_years"],
            **orbit_computed,
        },
        "system_description": {
            "spacecraft_count": sum(s.get("quantity", 1) for s in spacecraft),
            "spacecraft": [{"name": s["name"], "quantity": s.get("quantity", 1)} for s in spacecraft],
            "total_mass_kg": round(total_mass, 2),
            "form_factor": form_factor,
        },
        "orbital_lifetime": {
            "orbital_lifetime_years": orbital_lifetime_years,
            "deorbit_compliant_25yr": deorbit_compliant_25yr,
        },
        "gsd_estimate_m": gsd_estimate_m,
        "ground_stations": gs_list,
        "frequency_usage": [
            {"subsystem": e["name"], "bands": (e.get("performance") or {}).get("bands", []),
             "rf_band": (e.get("performance") or {}).get("rf_band")}
            for e in ttc_elements
        ],
        "template_fields": filing_data,
    }

    if fmt.lower() == "docx":
        from ..services.branding import create_branded_docx
        doc = create_branded_docx("RSSSA Licence Application Data", f"Study {study_id}")
        if doc is None:
            raise HTTPException(status_code=500, detail="python-docx not available")

        doc.add_heading("1. Applicant Information", level=1)
        doc.add_paragraph(f"Name: {b.university}")
        doc.add_paragraph(f"Department: {b.department}")
        doc.add_paragraph(f"Country: Canada")

        doc.add_heading("2. Orbital Parameters", level=1)
        doc.add_paragraph(f"Altitude: {orbit_params['altitude_km']} km")
        doc.add_paragraph(f"Inclination: {orbit_params['inclination_deg']} deg")
        doc.add_paragraph(f"Orbit type: {orbit_params['orbit_type']}")
        doc.add_paragraph(f"Eccentricity: {orbit_params['eccentricity']}")
        doc.add_paragraph(f"Design lifetime: {orbit_params['design_lifetime_years']} years")
        if orbit_computed:
            doc.add_paragraph(f"Period: {orbit_computed.get('period_min', 'TBD')} min")
            doc.add_paragraph(f"Eclipse fraction: {orbit_computed.get('eclipse_fraction', 'TBD')}")
            doc.add_paragraph(f"Orbits per day: {orbit_computed.get('orbits_per_day', 'TBD')}")

        doc.add_heading("3. System Description", level=1)
        doc.add_paragraph(f"Spacecraft count: {data['system_description']['spacecraft_count']}")
        doc.add_paragraph(f"Total mass: {round(total_mass, 2)} kg")
        doc.add_paragraph(f"Form factor: {form_factor}")
        if spacecraft:
            t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Spacecraft"
            t.rows[0].cells[1].text = "Qty"
            for s in spacecraft:
                row = t.add_row()
                row.cells[0].text = s["name"]
                row.cells[1].text = str(s.get("quantity", 1))

        doc.add_heading("4. Orbital Lifetime Analysis", level=1)
        if orbital_lifetime_years is not None:
            doc.add_paragraph(f"Estimated orbital lifetime: {orbital_lifetime_years} years")
            doc.add_paragraph(f"25-year deorbit compliant: {'Yes' if deorbit_compliant_25yr else 'No'}")
        else:
            doc.add_paragraph("Orbital lifetime analysis not available (physics module missing).")

        if gsd_estimate_m is not None:
            doc.add_heading("5. Imaging Capability", level=1)
            doc.add_paragraph(f"Estimated GSD: {gsd_estimate_m} m")
            doc.add_paragraph(f"Pointing accuracy: {pointing_accuracy_deg} deg")

        next_section = 6 if gsd_estimate_m is not None else 5
        doc.add_heading(f"{next_section}. Ground Stations", level=1)
        if gs_elements:
            t = doc.add_table(rows=1, cols=4, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Station"
            t.rows[0].cells[1].text = "Latitude"
            t.rows[0].cells[2].text = "Longitude"
            t.rows[0].cells[3].text = "Bands"
            for e in gs_elements:
                row = t.add_row()
                row.cells[0].text = e["name"]
                row.cells[1].text = str(e["performance"]["latitude"])
                row.cells[2].text = str(e["performance"]["longitude"])
                row.cells[3].text = ", ".join(e["performance"].get("bands", []))
        else:
            doc.add_paragraph("No ground stations defined.")

        doc.add_heading(f"{next_section + 1}. Frequency Usage", level=1)
        if ttc_elements:
            t = doc.add_table(rows=1, cols=3, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Subsystem"
            t.rows[0].cells[1].text = "Bands"
            t.rows[0].cells[2].text = "RF Band"
            for e in ttc_elements:
                row = t.add_row()
                row.cells[0].text = e["name"]
                perf = e.get("performance") or {}
                row.cells[1].text = ", ".join(perf.get("bands", []))
                row.cells[2].text = str(perf.get("rf_band") or "TBD")
        else:
            doc.add_paragraph("No TT&C elements defined.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="RSSSA_{study_id}.docx"'},
        )

    return JSONResponse(content=data)


@router.post("/deorbit/{study_id}")
async def export_deorbit(
    study_id: str,
    fmt: str = Query(default="json", description="Output format: json, docx"),
) -> JSONResponse:
    """Generate deorbit analysis and debris compliance report."""
    from ..services.branding import get_branding
    from .elements import _elements

    elements = [e for e in _elements.values() if e.get("study_id") == study_id and not e.get("deleted_at")]
    total_mass = sum((e.get("mass_kg") or 0) * e.get("quantity", 1) for e in elements if e.get("segment") == "space")

    # Read actual orbit altitude from study requirements, fall back to 500 km
    altitude_km = 500.0
    studies = get_study_store()
    study = studies.get(study_id)
    orbit_params = _read_orbit_params(study) if study else {
        "altitude_km": 500.0, "inclination_deg": 97.4, "orbit_type": "SSO",
        "eccentricity": 0.0, "design_lifetime_years": 3.0,
    }
    altitude_km = orbit_params["altitude_km"]

    # Derive form factor and cross-section
    form_factor = _form_factor_from_mass(total_mass)
    cross_section_m2 = 0.03  # default
    try:
        from spacecdf_common.physics.debris import estimate_cubesat_cross_section
        if form_factor != "Custom":
            cross_section_m2 = estimate_cubesat_cross_section(form_factor)
    except Exception:
        pass

    a_over_m = cross_section_m2 / max(total_mass, 0.1) if total_mass > 0 else 0.01

    # Run deorbit analysis using existing physics — 3 solar scenarios
    deorbit_data: dict = {}
    lifetime_scenarios: list[dict] = []
    try:
        from spacecdf_common.physics.debris import (
            estimate_orbital_lifetime, compute_casualty_risk, check_deorbit_compliance
        )
        # 3 solar scenarios
        for label, f107 in [("Solar Minimum (F10.7=70)", 70.0),
                            ("Solar Mean (F10.7=120)", 120.0),
                            ("Solar Maximum (F10.7=200)", 200.0)]:
            lt = estimate_orbital_lifetime(
                altitude_km=altitude_km,
                area_to_mass_ratio_m2_kg=a_over_m,
                f107=f107,
            )
            lt_rounded = round(lt, 2) if isinstance(lt, (int, float)) and lt < 1e5 else None
            lifetime_scenarios.append({
                "scenario": label,
                "f107": f107,
                "lifetime_years": lt_rounded,
                "compliant_25yr": lt_rounded <= 25.0 if lt_rounded is not None else None,
                "compliant_5yr": lt_rounded <= 5.0 if lt_rounded is not None else None,
            })

        # Use mean scenario for primary result
        mean_lifetime = lifetime_scenarios[1]["lifetime_years"]
        casualty = compute_casualty_risk(mass_kg=total_mass or 6)
        deorbit_data = {
            "orbital_lifetime_years": mean_lifetime,
            "casualty_risk": round(casualty, 6) if isinstance(casualty, (int, float)) else None,
            "compliant_25yr": mean_lifetime <= 25.0 if mean_lifetime is not None else None,
            "compliant_5yr": mean_lifetime <= 5.0 if mean_lifetime is not None else None,
        }
    except Exception:
        # Fallback: try the older API
        try:
            from spacecdf_common.physics.debris import (
                compute_orbital_lifetime, compute_casualty_risk, check_deorbit_compliance
            )
            lifetime = compute_orbital_lifetime(altitude_km=altitude_km, mass_kg=total_mass or 6, area_m2=cross_section_m2)
            casualty = compute_casualty_risk(mass_kg=total_mass or 6)
            compliance = check_deorbit_compliance(altitude_km=altitude_km, mass_kg=total_mass or 6, area_m2=cross_section_m2)
            deorbit_data = {
                "orbital_lifetime_years": round(lifetime, 1) if isinstance(lifetime, (int, float)) else None,
                "casualty_risk": round(casualty, 6) if isinstance(casualty, (int, float)) else None,
                "compliant_25yr": compliance if isinstance(compliance, bool) else None,
            }
        except Exception:
            deorbit_data = {"note": "Deorbit physics module not available — manual analysis required"}

    # Build passivation checklist from element types
    passivation_checklist = []
    has_battery = any(
        e.get("subsystem_domain") in ("eps", "power") or
        "battery" in (e.get("name") or "").lower()
        for e in elements if e.get("segment") == "space"
    )
    has_propulsion = any(
        e.get("subsystem_domain") == "propulsion" or
        "thruster" in (e.get("name") or "").lower() or
        "propulsion" in (e.get("name") or "").lower()
        for e in elements if e.get("segment") == "space"
    )
    has_ttc = any(
        e.get("subsystem_domain") == "ttc" or
        "transmitter" in (e.get("name") or "").lower() or
        "radio" in (e.get("name") or "").lower()
        for e in elements if e.get("segment") == "space"
    )
    has_momentum_wheels = any(
        e.get("subsystem_domain") == "aocs" or
        "wheel" in (e.get("name") or "").lower() or
        "magnetorquer" in (e.get("name") or "").lower()
        for e in elements if e.get("segment") == "space"
    )

    if has_battery:
        passivation_checklist.append({
            "item": "Discharge batteries",
            "description": "Discharge batteries to safe level (< 50% SoC) to prevent rupture",
            "standard": "ISO 24113 Cl. 6.2.3.1",
        })
    if has_propulsion:
        passivation_checklist.append({
            "item": "Vent propellant",
            "description": "Deplete or vent remaining propellant and pressurant",
            "standard": "ISO 24113 Cl. 6.2.3.2",
        })
    if has_ttc:
        passivation_checklist.append({
            "item": "Disable transmitter",
            "description": "Permanently disable RF transmitter to free frequency allocation",
            "standard": "ITU Radio Regulations Art. 22",
        })
    if has_momentum_wheels:
        passivation_checklist.append({
            "item": "Spin down momentum wheels",
            "description": "De-spin reaction/momentum wheels to prevent debris generation",
            "standard": "ISO 24113 Cl. 6.2.3.3",
        })
    # Always include solar array isolation
    passivation_checklist.append({
        "item": "Isolate solar arrays",
        "description": "Open solar array circuit to prevent battery charging",
        "standard": "ISO 24113 Cl. 6.2.3.1",
    })

    # Compliance summary
    compliance_summary = {
        "iadc_25yr": deorbit_data.get("compliant_25yr"),
        "fcc_5yr": deorbit_data.get("compliant_5yr"),
        "casualty_risk_compliant": (
            deorbit_data.get("casualty_risk", 1.0) < 0.0001
            if isinstance(deorbit_data.get("casualty_risk"), (int, float)) else None
        ),
        "passivation_items_count": len(passivation_checklist),
    }

    b = get_branding()
    data = {
        "document": "Deorbit Analysis & Debris Compliance Report",
        "branding": {"university": b.university, "department": b.department, "classification": b.classification},
        "study_id": study_id,
        "spacecraft": {
            "total_mass_kg": round(total_mass, 2),
            "altitude_km": altitude_km,
            "form_factor": form_factor,
            "cross_section_m2": round(cross_section_m2, 4),
        },
        "analysis": deorbit_data,
        "lifetime_scenarios": lifetime_scenarios,
        "passivation_checklist": passivation_checklist,
        "compliance_summary": compliance_summary,
        "standards": {
            "iso_24113": "ISO 24113:2023 — Space debris mitigation requirements",
            "ecss_u_as_10c": "ECSS-U-AS-10C Rev.2 — Space sustainability",
            "fcc_5yr": "FCC 2024+ 5-year deorbit rule",
            "iadc_25yr": "IADC 25-year guideline",
        },
        "mitigation_options": [
            {"method": "Natural decay", "description": "Rely on atmospheric drag at current altitude"},
            {"method": "Propulsive deorbit", "description": "Use onboard thruster for controlled reentry"},
            {"method": "Drag sail", "description": "Deploy drag augmentation device at end-of-life"},
            {"method": "Electrodynamic tether", "description": "Lorentz force deorbiting"},
        ],
    }

    if fmt.lower() == "docx":
        from ..services.branding import create_branded_docx
        doc = create_branded_docx("Deorbit Analysis & Debris Compliance Report", f"Study {study_id}")
        if doc is None:
            raise HTTPException(status_code=500, detail="python-docx not available")

        doc.add_heading("1. Spacecraft Parameters", level=1)
        doc.add_paragraph(f"Total mass: {round(total_mass, 2)} kg")
        doc.add_paragraph(f"Orbit altitude: {altitude_km} km")
        doc.add_paragraph(f"Form factor: {form_factor}")
        doc.add_paragraph(f"Cross-sectional area: {round(cross_section_m2, 4)} m\u00b2")

        doc.add_heading("2. Lifetime Analysis — 3 Solar Scenarios", level=1)
        if lifetime_scenarios:
            t = doc.add_table(rows=1, cols=4, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Scenario"
            t.rows[0].cells[1].text = "F10.7"
            t.rows[0].cells[2].text = "Lifetime (yr)"
            t.rows[0].cells[3].text = "25-yr Compliant"
            for sc in lifetime_scenarios:
                row = t.add_row()
                row.cells[0].text = sc["scenario"]
                row.cells[1].text = str(sc["f107"])
                row.cells[2].text = str(sc["lifetime_years"]) if sc["lifetime_years"] is not None else "> 100,000"
                row.cells[3].text = "Yes" if sc.get("compliant_25yr") else "No" if sc.get("compliant_25yr") is False else "N/A"
        elif "note" in deorbit_data:
            doc.add_paragraph(deorbit_data["note"])
        else:
            t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Parameter"
            t.rows[0].cells[1].text = "Value"
            for key, val in deorbit_data.items():
                row = t.add_row()
                row.cells[0].text = key.replace("_", " ").title()
                row.cells[1].text = str(val)

        doc.add_heading("3. Passivation Checklist", level=1)
        if passivation_checklist:
            t = doc.add_table(rows=1, cols=3, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Action"
            t.rows[0].cells[1].text = "Description"
            t.rows[0].cells[2].text = "Standard"
            for item in passivation_checklist:
                row = t.add_row()
                row.cells[0].text = item["item"]
                row.cells[1].text = item["description"]
                row.cells[2].text = item["standard"]
        else:
            doc.add_paragraph("No passivation items identified.")

        doc.add_heading("4. Compliance Summary", level=1)
        t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
        t.rows[0].cells[0].text = "Requirement"
        t.rows[0].cells[1].text = "Status"
        for label, key in [("IADC 25-year guideline", "iadc_25yr"),
                           ("FCC 5-year rule", "fcc_5yr"),
                           ("Casualty risk < 1:10000", "casualty_risk_compliant")]:
            row = t.add_row()
            row.cells[0].text = label
            val = compliance_summary.get(key)
            row.cells[1].text = "PASS" if val else "FAIL" if val is False else "N/A"

        doc.add_heading("5. Applicable Standards", level=1)
        for key, std in data["standards"].items():
            doc.add_paragraph(std, style="List Bullet")

        doc.add_heading("6. Mitigation Options", level=1)
        t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
        t.rows[0].cells[0].text = "Method"
        t.rows[0].cells[1].text = "Description"
        for opt in data["mitigation_options"]:
            row = t.add_row()
            row.cells[0].text = opt["method"]
            row.cells[1].text = opt["description"]

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Deorbit_Report_{study_id}.docx"'},
        )

    return JSONResponse(content=data)


@router.post("/thermal-report/{study_id}")
async def export_thermal_report(
    study_id: str,
    fmt: str = Query(default="json", description="Output format: json, docx"),
) -> JSONResponse:
    """Generate thermal design report data."""
    from ..services.branding import get_branding
    from .elements import _elements

    # Read study for orbit params
    studies = get_study_store()
    study = studies.get(study_id)
    orbit_params = _read_orbit_params(study) if study else {
        "altitude_km": 500.0, "inclination_deg": 97.4, "orbit_type": "SSO",
        "eccentricity": 0.0, "design_lifetime_years": 3.0,
    }

    elements = [e for e in _elements.values() if e.get("study_id") == study_id and not e.get("deleted_at")]
    space_elements = [e for e in elements if e.get("segment") == "space"]
    power_elements = [e for e in elements if (e.get("power_avg_w") or 0) > 0]

    total_power = sum((e.get("power_avg_w") or 0) * e.get("quantity", 1) for e in power_elements)
    total_mass = sum((e.get("mass_kg") or 0) * e.get("quantity", 1) for e in space_elements)
    form_factor = _form_factor_from_mass(total_mass)

    # Compute eclipse fraction from orbit params
    eclipse_fraction = 0.35  # default
    orbit_computed: dict = {}
    try:
        from spacecdf_common.physics.orbit import compute_orbit_params
        op = compute_orbit_params(
            altitude_km=orbit_params["altitude_km"],
            inclination_deg=orbit_params["inclination_deg"],
            eccentricity=orbit_params["eccentricity"],
        )
        eclipse_fraction = op.eclipse_fraction
        orbit_computed = {
            "altitude_km": orbit_params["altitude_km"],
            "inclination_deg": orbit_params["inclination_deg"],
            "orbit_type": orbit_params["orbit_type"],
            "period_min": round(op.period_min, 2),
            "eclipse_fraction": round(op.eclipse_fraction, 3),
            "eclipse_duration_min": round(op.eclipse_duration_min, 2),
            "sunlight_duration_min": round(op.sunlight_duration_min, 2),
        }
    except Exception:
        orbit_computed = {
            "altitude_km": orbit_params["altitude_km"],
            "inclination_deg": orbit_params["inclination_deg"],
            "orbit_type": orbit_params["orbit_type"],
        }

    # Estimate spacecraft surface area
    spacecraft_area_m2 = 1.0
    try:
        from spacecdf_common.physics.thermal import spacecraft_surface_area
        spacecraft_area_m2 = spacecraft_surface_area(total_mass, form_factor="cubesat" if form_factor != "Custom" else "box")
    except Exception:
        pass

    # Run thermal balance computation
    thermal_result: dict = {}
    try:
        from spacecdf_common.physics.thermal import compute_thermal_balance
        # Read temperature limits from study requirements if available
        t_min_c = -20.0
        t_max_c = 50.0
        try:
            temp_range = study.requirements.payloads[0].temperature_range_c
            if temp_range and len(temp_range) >= 2:
                t_min_c = temp_range[0]
                t_max_c = temp_range[1]
        except Exception:
            pass

        tb = compute_thermal_balance(
            internal_power_w=total_power if total_power > 0 else 10.0,
            spacecraft_area_m2=spacecraft_area_m2,
            eclipse_fraction=eclipse_fraction,
            t_min_c=t_min_c,
            t_max_c=t_max_c,
        )
        thermal_result = {
            "hot_case_temp_c": round(tb.hot_case_temp_c, 1),
            "cold_case_temp_c": round(tb.cold_case_temp_c, 1),
            "cold_case_heater_w": round(tb.cold_case_heater_w, 2),
            "radiator_area_m2": round(tb.radiator_area_m2, 4),
            "radiator_mass_kg": round(tb.radiator_mass_kg, 3),
            "mli_area_m2": round(tb.mli_area_m2, 4),
            "mli_mass_kg": round(tb.mli_mass_kg, 3),
            "tcs_mass_kg": round(tb.tcs_mass_kg, 3),
            "tcs_heater_power_w": round(tb.tcs_heater_power_w, 2),
            "warnings": tb.warnings,
        }
    except Exception:
        thermal_result = {"note": "Thermal balance computation not available"}

    # Build component temperature range table
    component_thermal = []
    for e in space_elements:
        if e.get("element_type") == "component":
            comp_entry = {
                "name": e["name"],
                "subsystem": e.get("subsystem_domain", ""),
                "power_avg_w": e.get("power_avg_w", 0),
                "quantity": e.get("quantity", 1),
            }
            # Look for component-level temp limits in performance dict
            perf = e.get("performance") or {}
            if perf.get("temperature_range_c"):
                comp_entry["temp_range_c"] = perf["temperature_range_c"]
            elif perf.get("operating_temp_min_c") is not None:
                comp_entry["temp_range_c"] = [perf["operating_temp_min_c"], perf.get("operating_temp_max_c", 50)]
            else:
                # Use typical ranges by subsystem domain
                domain = e.get("subsystem_domain", "")
                typical_ranges = {
                    "eps": [-20, 50], "payload": [-10, 40], "obc": [-20, 60],
                    "ttc": [-20, 50], "aocs": [-20, 50], "propulsion": [5, 50],
                    "structure": [-40, 80], "thermal": [-40, 80],
                }
                comp_entry["temp_range_c"] = typical_ranges.get(domain, [-20, 50])
            component_thermal.append(comp_entry)

    b = get_branding()
    data = {
        "document": "Thermal Design Report",
        "branding": {"university": b.university, "department": b.department, "classification": b.classification},
        "study_id": study_id,
        "orbit": orbit_computed,
        "thermal_environment": {
            "orbit": f"{orbit_params['orbit_type']} ({orbit_params['altitude_km']} km)",
            "eclipse_fraction": round(eclipse_fraction, 3),
            "solar_flux_w_m2": 1361,
            "albedo_factor": 0.3,
            "earth_ir_w_m2": 237,
            "spacecraft_area_m2": round(spacecraft_area_m2, 4),
        },
        "power_dissipation": {
            "total_avg_w": round(total_power, 1),
            "elements": [
                {"name": e["name"], "power_avg_w": e.get("power_avg_w", 0), "quantity": e.get("quantity", 1),
                 "domain": e.get("subsystem_domain", "")}
                for e in power_elements
            ],
        },
        "thermal_balance": thermal_result,
        "component_thermal_ranges": component_thermal,
        "design_notes": [
            "Passive thermal control assumed (surface coatings + MLI)",
            "Heater power allocated for eclipse survival",
            "Radiator area sized from total power dissipation",
        ],
    }

    if fmt.lower() == "docx":
        from ..services.branding import create_branded_docx
        doc = create_branded_docx("Thermal Design Report", f"Study {study_id}")
        if doc is None:
            raise HTTPException(status_code=500, detail="python-docx not available")

        doc.add_heading("1. Orbital Environment", level=1)
        doc.add_paragraph(f"Orbit: {orbit_params['orbit_type']} at {orbit_params['altitude_km']} km")
        doc.add_paragraph(f"Eclipse fraction: {round(eclipse_fraction, 3)}")
        if orbit_computed.get("period_min"):
            doc.add_paragraph(f"Period: {orbit_computed['period_min']} min")
            doc.add_paragraph(f"Eclipse duration: {orbit_computed.get('eclipse_duration_min', 'TBD')} min")
            doc.add_paragraph(f"Sunlight duration: {orbit_computed.get('sunlight_duration_min', 'TBD')} min")

        doc.add_heading("2. Thermal Environment", level=1)
        te = data["thermal_environment"]
        doc.add_paragraph(f"Solar flux: {te['solar_flux_w_m2']} W/m\u00b2")
        doc.add_paragraph(f"Albedo factor: {te['albedo_factor']}")
        doc.add_paragraph(f"Earth IR: {te['earth_ir_w_m2']} W/m\u00b2")
        doc.add_paragraph(f"Spacecraft surface area: {te['spacecraft_area_m2']} m\u00b2")

        doc.add_heading("3. Power Dissipation", level=1)
        doc.add_paragraph(f"Total average power: {round(total_power, 1)} W")
        if power_elements:
            t = doc.add_table(rows=1, cols=4, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Element"
            t.rows[0].cells[1].text = "Power (W)"
            t.rows[0].cells[2].text = "Qty"
            t.rows[0].cells[3].text = "Domain"
            for e in power_elements:
                row = t.add_row()
                row.cells[0].text = e["name"]
                row.cells[1].text = str(e.get("power_avg_w", 0))
                row.cells[2].text = str(e.get("quantity", 1))
                row.cells[3].text = e.get("subsystem_domain", "")

        doc.add_heading("4. Thermal Balance Analysis", level=1)
        if "note" in thermal_result:
            doc.add_paragraph(thermal_result["note"])
        else:
            t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Parameter"
            t.rows[0].cells[1].text = "Value"
            for label, key, unit in [
                ("Hot case temperature", "hot_case_temp_c", "C"),
                ("Cold case temperature", "cold_case_temp_c", "C"),
                ("Eclipse heater power", "cold_case_heater_w", "W"),
                ("Radiator area", "radiator_area_m2", "m\u00b2"),
                ("Radiator mass", "radiator_mass_kg", "kg"),
                ("MLI area", "mli_area_m2", "m\u00b2"),
                ("MLI mass", "mli_mass_kg", "kg"),
                ("TCS total mass", "tcs_mass_kg", "kg"),
                ("TCS heater power", "tcs_heater_power_w", "W"),
            ]:
                if key in thermal_result:
                    row = t.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = f"{thermal_result[key]} {unit}"
            if thermal_result.get("warnings"):
                doc.add_paragraph("")
                for w in thermal_result["warnings"]:
                    doc.add_paragraph(f"WARNING: {w}", style="List Bullet")

        doc.add_heading("5. Component Temperature Ranges", level=1)
        if component_thermal:
            t = doc.add_table(rows=1, cols=5, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Component"
            t.rows[0].cells[1].text = "Subsystem"
            t.rows[0].cells[2].text = "Power (W)"
            t.rows[0].cells[3].text = "T_min (C)"
            t.rows[0].cells[4].text = "T_max (C)"
            for c in component_thermal:
                row = t.add_row()
                row.cells[0].text = c["name"]
                row.cells[1].text = c.get("subsystem", "")
                row.cells[2].text = str(c.get("power_avg_w", 0))
                tr = c.get("temp_range_c", [-20, 50])
                row.cells[3].text = str(tr[0])
                row.cells[4].text = str(tr[1])
        else:
            doc.add_paragraph("No space segment components defined.")

        doc.add_heading("6. Design Notes", level=1)
        for note in data["design_notes"]:
            doc.add_paragraph(note, style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Thermal_Report_{study_id}.docx"'},
        )

    return JSONResponse(content=data)


@router.post("/test-plan/{study_id}")
async def export_test_plan(
    study_id: str,
    fmt: str = Query(default="json", description="Output format: json, docx, xlsx"),
) -> JSONResponse:
    """Generate AIT/AIV test plan from requirements."""
    from ..services.branding import get_branding
    from ..routers.requirements import _requirements
    from .elements import _elements, _interfaces

    # Read study
    studies = get_study_store()
    study = studies.get(study_id)

    reqs = [r for r in _requirements.values() if r.get("study_id") == study_id and r.get("status") != "retired"]
    test_reqs = [r for r in reqs if r.get("verification_method") == "T"]
    analysis_reqs = [r for r in reqs if r.get("verification_method") == "A"]
    inspection_reqs = [r for r in reqs if r.get("verification_method") == "I"]

    # Read elements for component-level tests
    elements = [e for e in _elements.values() if e.get("study_id") == study_id and not e.get("deleted_at")]
    space_components = [e for e in elements if e.get("segment") == "space" and e.get("element_type") == "component"]

    # Read interfaces for integration test matrix
    study_interfaces = [
        i for i in _interfaces.values()
        if not i.get("deleted_at") and (
            i.get("study_id") == study_id or
            i.get("source_element_id") in {e.get("id") for e in elements} or
            i.get("target_element_id") in {e.get("id") for e in elements}
        )
    ]

    # Build element ID -> name lookup
    element_names = {e.get("id"): e.get("name", "Unknown") for e in elements}

    test_cases = []
    for i, r in enumerate(test_reqs):
        test_cases.append({
            "test_id": f"TC-{i+1:03d}",
            "requirement_code": r.get("code", r["id"]),
            "requirement_text": r.get("text", ""),
            "level": r.get("level", "system"),
            "test_description": f"Verify: {r.get('text', '')}",
            "pass_criteria": f"Requirement {r.get('code', '')} is satisfied",
            "test_type": "Functional",
            "status": "planned",
        })

    # Component-level unit tests (one per component)
    unit_tests = []
    for j, comp in enumerate(space_components):
        unit_tests.append({
            "test_id": f"UT-{j+1:03d}",
            "component": comp["name"],
            "subsystem": comp.get("subsystem_domain", ""),
            "test_description": f"Functional verification of {comp['name']}",
            "pass_criteria": f"{comp['name']} operates within specified parameters",
            "checks": [],
        })
        # Add specific checks based on component properties
        if comp.get("mass_kg"):
            unit_tests[-1]["checks"].append(f"Mass verification: {comp['mass_kg']} kg +/- 5%")
        if comp.get("power_avg_w"):
            unit_tests[-1]["checks"].append(f"Power draw: {comp['power_avg_w']} W nominal")
        perf = comp.get("performance") or {}
        if perf.get("data_rate_mbps"):
            unit_tests[-1]["checks"].append(f"Data rate: {perf['data_rate_mbps']} Mbps")
        if perf.get("pointing_accuracy_deg"):
            unit_tests[-1]["checks"].append(f"Pointing accuracy: {perf['pointing_accuracy_deg']} deg")

    # Integration test matrix from interfaces
    integration_tests = []
    for k, iface in enumerate(study_interfaces):
        src_name = element_names.get(iface.get("source_element_id"), iface.get("source_name", "Unknown"))
        tgt_name = element_names.get(iface.get("target_element_id"), iface.get("target_name", "Unknown"))
        iface_type = iface.get("interface_type", iface.get("type", "data"))
        integration_tests.append({
            "test_id": f"IT-{k+1:03d}",
            "source": src_name,
            "target": tgt_name,
            "interface_type": iface_type,
            "test_description": f"Verify {iface_type} interface between {src_name} and {tgt_name}",
            "pass_criteria": f"Data/power/signal flows correctly between {src_name} and {tgt_name}",
        })

    # Environmental test profile from launch vehicle
    env_profile: dict = {}
    try:
        from spacecdf_common.physics.structures import launch_loads
        # Try to detect launch vehicle from study notes or elements
        launch_vehicle = "falcon_9"  # default
        if study:
            notes = getattr(study, "notes", "") or ""
            for lv in ["electron", "falcon_9", "falcon_heavy", "vega_c", "ariane_6",
                        "pslv", "soyuz", "h3", "sls", "starship", "new_glenn", "long_march_5"]:
                if lv.replace("_", " ") in notes.lower() or lv in notes.lower():
                    launch_vehicle = lv
                    break

        axial_g, lateral_g = launch_loads(launch_vehicle)
        env_profile = {
            "launch_vehicle": launch_vehicle,
            "quasi_static_loads": {
                "axial_g": axial_g,
                "lateral_g": lateral_g,
            },
            "random_vibration": {
                "level_grms": round(axial_g * 2.35, 1),  # Typical qualification = ~2.35x QSL
                "frequency_range_hz": "20 - 2000",
                "duration_s": 120,
            },
            "sine_vibration": {
                "axial_g": round(axial_g * 1.25, 1),
                "lateral_g": round(lateral_g * 1.25, 1),
                "sweep_rate_oct_min": 2,
                "frequency_range_hz": "5 - 100",
            },
            "shock": {
                "srs_g_at_1000hz": 1500,
                "srs_g_at_3000hz": 3000,
            },
            "thermal_vacuum": {
                "hot_survival_c": 60,
                "cold_survival_c": -40,
                "hot_operational_c": 50,
                "cold_operational_c": -20,
                "cycles": 4,
                "pressure_mbar": 1e-5,
            },
        }
    except Exception:
        env_profile = {
            "note": "Launch loads module not available — using generic profile",
            "quasi_static_loads": {"axial_g": 6.0, "lateral_g": 2.0},
            "thermal_vacuum": {
                "hot_survival_c": 60, "cold_survival_c": -40,
                "hot_operational_c": 50, "cold_operational_c": -20,
                "cycles": 4,
            },
        }

    b = get_branding()
    data = {
        "document": "AIT/AIV Test Plan",
        "branding": {"university": b.university, "department": b.department, "classification": b.classification},
        "study_id": study_id,
        "summary": {
            "total_requirements": len(reqs),
            "test_requirements": len(test_reqs),
            "analysis_requirements": len(analysis_reqs),
            "inspection_requirements": len(inspection_reqs),
            "total_components": len(space_components),
            "total_interfaces": len(study_interfaces),
            "unit_tests": len(unit_tests),
            "integration_tests": len(integration_tests),
        },
        "test_cases": test_cases,
        "unit_tests": unit_tests,
        "integration_tests": integration_tests,
        "environmental_profile": env_profile,
        "test_phases": [
            {"phase": "Unit Test", "description": "Individual component functional verification"},
            {"phase": "Integration Test", "description": "Subsystem-level interface and performance verification"},
            {"phase": "System Test", "description": "Full spacecraft functional and environmental testing"},
            {"phase": "Acceptance Test", "description": "Final verification before launch campaign"},
        ],
    }

    if fmt.lower() == "docx":
        from ..services.branding import create_branded_docx
        doc = create_branded_docx("AIT/AIV Test Plan", f"Study {study_id}")
        if doc is None:
            raise HTTPException(status_code=500, detail="python-docx not available")

        doc.add_heading("1. Summary", level=1)
        doc.add_paragraph(f"Total requirements: {len(reqs)}")
        doc.add_paragraph(f"Test (T): {len(test_reqs)}  |  Analysis (A): {len(analysis_reqs)}  |  Inspection (I): {len(inspection_reqs)}")
        doc.add_paragraph(f"Components: {len(space_components)}  |  Interfaces: {len(study_interfaces)}")
        doc.add_paragraph(f"Unit tests: {len(unit_tests)}  |  Integration tests: {len(integration_tests)}")

        doc.add_heading("2. Requirement-Level Test Cases", level=1)
        if test_cases:
            t = doc.add_table(rows=1, cols=6, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Test ID"
            t.rows[0].cells[1].text = "Req Code"
            t.rows[0].cells[2].text = "Level"
            t.rows[0].cells[3].text = "Description"
            t.rows[0].cells[4].text = "Pass Criteria"
            t.rows[0].cells[5].text = "Status"
            for tc in test_cases:
                row = t.add_row()
                row.cells[0].text = tc["test_id"]
                row.cells[1].text = tc["requirement_code"]
                row.cells[2].text = tc["level"]
                row.cells[3].text = tc["test_description"]
                row.cells[4].text = tc["pass_criteria"]
                row.cells[5].text = tc["status"]
        else:
            doc.add_paragraph("No test-verified requirements defined.")

        doc.add_heading("3. Unit Tests (per Component)", level=1)
        if unit_tests:
            for ut in unit_tests:
                doc.add_heading(f"{ut['test_id']}: {ut['component']}", level=2)
                doc.add_paragraph(f"Subsystem: {ut['subsystem']}")
                doc.add_paragraph(f"Description: {ut['test_description']}")
                doc.add_paragraph(f"Pass criteria: {ut['pass_criteria']}")
                if ut["checks"]:
                    doc.add_paragraph("Specific checks:")
                    for chk in ut["checks"]:
                        doc.add_paragraph(chk, style="List Bullet")
        else:
            doc.add_paragraph("No space segment components defined for unit testing.")

        doc.add_heading("4. Integration Test Matrix", level=1)
        if integration_tests:
            t = doc.add_table(rows=1, cols=5, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Test ID"
            t.rows[0].cells[1].text = "Source"
            t.rows[0].cells[2].text = "Target"
            t.rows[0].cells[3].text = "Interface Type"
            t.rows[0].cells[4].text = "Description"
            for it in integration_tests:
                row = t.add_row()
                row.cells[0].text = it["test_id"]
                row.cells[1].text = it["source"]
                row.cells[2].text = it["target"]
                row.cells[3].text = it["interface_type"]
                row.cells[4].text = it["test_description"]
        else:
            doc.add_paragraph("No interfaces defined for integration testing.")

        doc.add_heading("5. Environmental Test Profile", level=1)
        if env_profile.get("note"):
            doc.add_paragraph(env_profile["note"])
        if env_profile.get("launch_vehicle"):
            doc.add_paragraph(f"Launch vehicle: {env_profile['launch_vehicle']}")

        qsl = env_profile.get("quasi_static_loads", {})
        if qsl:
            doc.add_heading("Quasi-Static Loads", level=2)
            doc.add_paragraph(f"Axial: {qsl.get('axial_g', 'TBD')} g")
            doc.add_paragraph(f"Lateral: {qsl.get('lateral_g', 'TBD')} g")

        rv = env_profile.get("random_vibration")
        if rv:
            doc.add_heading("Random Vibration", level=2)
            doc.add_paragraph(f"Level: {rv.get('level_grms', 'TBD')} grms")
            doc.add_paragraph(f"Frequency range: {rv.get('frequency_range_hz', 'TBD')}")
            doc.add_paragraph(f"Duration: {rv.get('duration_s', 'TBD')} s per axis")

        sv = env_profile.get("sine_vibration")
        if sv:
            doc.add_heading("Sine Vibration", level=2)
            doc.add_paragraph(f"Axial: {sv.get('axial_g', 'TBD')} g")
            doc.add_paragraph(f"Lateral: {sv.get('lateral_g', 'TBD')} g")
            doc.add_paragraph(f"Sweep rate: {sv.get('sweep_rate_oct_min', 'TBD')} oct/min")

        shock = env_profile.get("shock")
        if shock:
            doc.add_heading("Shock (SRS)", level=2)
            doc.add_paragraph(f"SRS @ 1000 Hz: {shock.get('srs_g_at_1000hz', 'TBD')} g")
            doc.add_paragraph(f"SRS @ 3000 Hz: {shock.get('srs_g_at_3000hz', 'TBD')} g")

        tvac = env_profile.get("thermal_vacuum")
        if tvac:
            doc.add_heading("Thermal Vacuum", level=2)
            t = doc.add_table(rows=1, cols=2, style="Light List Accent 1")
            t.rows[0].cells[0].text = "Parameter"
            t.rows[0].cells[1].text = "Value"
            for label, key, unit in [
                ("Hot survival", "hot_survival_c", "C"),
                ("Cold survival", "cold_survival_c", "C"),
                ("Hot operational", "hot_operational_c", "C"),
                ("Cold operational", "cold_operational_c", "C"),
                ("Cycles", "cycles", ""),
                ("Pressure", "pressure_mbar", "mbar"),
            ]:
                if key in tvac:
                    row = t.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = f"{tvac[key]} {unit}".strip()

        doc.add_heading("6. Test Phases", level=1)
        for phase in data["test_phases"]:
            doc.add_paragraph(f"{phase['phase']}: {phase['description']}", style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="Test_Plan_{study_id}.docx"'},
        )

    if fmt.lower() == "xlsx":
        from ..services.branding import create_branded_xlsx
        wb = create_branded_xlsx(f"AIT/AIV Test Plan — Study {study_id}")
        if wb is None:
            raise HTTPException(status_code=500, detail="openpyxl not available")

        ws = wb.create_sheet("Test Cases")
        headers = ["Test ID", "Req Code", "Req Text", "Level", "Description", "Pass Criteria", "Type", "Status"]
        ws.append(headers)
        from openpyxl.styles import Font as XlFont
        for cell in ws[1]:
            cell.font = XlFont(bold=True)
        for tc in test_cases:
            ws.append([
                tc["test_id"], tc["requirement_code"], tc["requirement_text"],
                tc["level"], tc["test_description"], tc["pass_criteria"],
                tc["test_type"], tc["status"],
            ])
        for col in ws.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        # Unit tests sheet
        ws_ut = wb.create_sheet("Unit Tests")
        ws_ut.append(["Test ID", "Component", "Subsystem", "Description", "Pass Criteria", "Checks"])
        for cell in ws_ut[1]:
            cell.font = XlFont(bold=True)
        for ut in unit_tests:
            ws_ut.append([
                ut["test_id"], ut["component"], ut["subsystem"],
                ut["test_description"], ut["pass_criteria"],
                "; ".join(ut["checks"]),
            ])
        for col in ws_ut.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws_ut.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        # Integration tests sheet
        ws_it = wb.create_sheet("Integration Tests")
        ws_it.append(["Test ID", "Source", "Target", "Interface Type", "Description", "Pass Criteria"])
        for cell in ws_it[1]:
            cell.font = XlFont(bold=True)
        for it in integration_tests:
            ws_it.append([
                it["test_id"], it["source"], it["target"],
                it["interface_type"], it["test_description"], it["pass_criteria"],
            ])
        for col in ws_it.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws_it.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        ws2 = wb.create_sheet("Summary")
        ws2.append(["Metric", "Count"])
        ws2.append(["Total requirements", len(reqs)])
        ws2.append(["Test (T)", len(test_reqs)])
        ws2.append(["Analysis (A)", len(analysis_reqs)])
        ws2.append(["Inspection (I)", len(inspection_reqs)])
        ws2.append(["Components", len(space_components)])
        ws2.append(["Interfaces", len(study_interfaces)])
        ws2.append(["Unit tests", len(unit_tests)])
        ws2.append(["Integration tests", len(integration_tests)])

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="Test_Plan_{study_id}.xlsx"'},
        )

    return JSONResponse(content=data)


# ─── Branded DOCX Worksheets ───

WORKSHEET_TEMPLATES = {
    "mission_need": {
        "title": "Mission Need Statement Worksheet",
        "sections": [
            ("1. Problem Statement", "What problem does this mission solve? Who is affected?"),
            ("2. Operational Context", "When, where, and how will mission products be used?"),
            ("3. Stakeholder Register", "Name / Role / Key Needs / Priority"),
            ("4. Mission Objectives", "ID / Objective / Priority / Measurable Criterion"),
            ("5. Success Criteria", "What constitutes mission success?"),
        ],
    },
    "trade_study": {
        "title": "Trade Study Worksheet",
        "sections": [
            ("1. Trade Context", "What decision is being made? What are the constraints?"),
            ("2. Evaluation Criteria", "Criterion / Weight / Direction / Threshold"),
            ("3. Options Under Consideration", "Option / Description / Pros / Cons"),
            ("4. Scoring Matrix", "Use Pugh Matrix or Weighted Scoring in Decide tab."),
            ("5. Decision and Rationale", "Selected option / Rationale / Dissenting opinions"),
        ],
    },
    "requirements": {
        "title": "Requirements Derivation Worksheet",
        "sections": [
            ("1. Parent Requirement", "Code / Text / Level"),
            ("2. Derived Requirements", "Code / Level / Text / Verification Method / Element"),
            ("3. Traceability Check", "All derived trace to parent? All parent covered?"),
            ("4. SMART Validation", "Specific? Measurable? Achievable? Relevant? Traceable?"),
        ],
    },
    "subsystem_design": {
        "title": "Subsystem Design Worksheet",
        "sections": [
            ("1. Subsystem Identification", "Name / Domain / Parent system / Driving requirement"),
            ("2. Architecture Options", "Option / Mass / Power / Cost / TRL / Notes"),
            ("3. Selected Architecture", "Selection / Rationale"),
            ("4. Budget Allocation", "Mass / Power / Cost allocations"),
            ("5. Interfaces", "Interface / Type / To-From / Direction / Properties"),
        ],
    },
    "equipment_selection": {
        "title": "Equipment Selection Worksheet",
        "sections": [
            ("1. Subsystem and Budget", "Subsystem / Mass budget / Power budget / Cost budget"),
            ("2. Candidate Components", "Component / Manufacturer / Mass / Power / Cost / TRL"),
            ("3. Selection and Rationale", "Selected / Rationale / Margin remaining"),
            ("4. Verification", "Inspection / Analysis / Test / Demonstration"),
        ],
    },
}


@router.post("/worksheet/{template_id}")
async def generate_worksheet(template_id: str) -> StreamingResponse:
    """Generate a branded DOCX worksheet template."""
    if template_id not in WORKSHEET_TEMPLATES:
        raise HTTPException(400, f"Unknown worksheet: {template_id}. Available: {list(WORKSHEET_TEMPLATES.keys())}")

    from ..services.branding import create_branded_docx
    from docx.shared import Pt

    tmpl = WORKSHEET_TEMPLATES[template_id]
    doc = create_branded_docx(tmpl["title"], "Worksheet Template")
    if doc is None:
        raise HTTPException(500, "python-docx not available")

    for heading, content in tmpl["sections"]:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
        # Add blank lines for writing
        for _ in range(3):
            p = doc.add_paragraph("")
            p.style.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="SpaceCDF_Worksheet_{template_id}.docx"'},
    )
